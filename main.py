#!/usr/bin/env python3

import json
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
import tkinterdnd2 as tkdnd
import os
import sys
import shutil
import sqlite3
import threading
from pathlib import Path
from src.pdf_parser import PDFParser
from src.name_declension import NameDeclension
from src.database import Database
from src.word_generator import WordGenerator
from src.config_loader import config_loader
from src.constants import OUTPUT_DIR, DATABASE_DIR, TEMPLATES_DIR
from src.data_mos_ru_sync import sync_organizations_from_data_mos_ru
from src.credential_manager import CredentialManager


def get_resource_path(relative_path):
    """
    Get absolute path to resource (works in dev mode and in .exe).

    Args:
        relative_path: Relative path to the resource

    Returns:
        Absolute path to the resource
    """
    try:
        base_path = sys._MEIPASS
    except AttributeError:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


def setup_application():
    """
    Automatic creation of necessary folders and copying templates on first run.
    Creates output, database, and templates directories if they don't exist.
    Extracts bundled templates from .exe if running as standalone executable.
    """
    output_dir = OUTPUT_DIR
    database_dir = DATABASE_DIR
    templates_dir = TEMPLATES_DIR
    config_dir = Path("config")

    output_dir.mkdir(exist_ok=True)
    database_dir.mkdir(exist_ok=True)

    bundled_templates = get_resource_path("templates")
    bundled_config = get_resource_path("config")

    if not templates_dir.exists() and os.path.exists(bundled_templates):
        print(f"[INFO] Extracting templates from .exe to {templates_dir}...")
        shutil.copytree(bundled_templates, templates_dir)
        print(f"[OK] Templates extracted to {templates_dir}")

    if not config_dir.exists() and os.path.exists(bundled_config):
        print(f"[INFO] Extracting config from .exe to {config_dir}...")
        shutil.copytree(bundled_config, config_dir)
        print(f"[OK] Config extracted to {config_dir}")

    if not templates_dir.exists():
        templates_dir.mkdir(exist_ok=True)
        print(f"[WARNING] WARNING: templates folder is empty!")
        print(f"[WARNING] Add .docx files to templates folder: {templates_dir.absolute()}")

    if not config_dir.exists():
        config_dir.mkdir(exist_ok=True)
        print(f"[WARNING] WARNING: config folder is empty!")
        print(f"[WARNING] Application may not work correctly without configuration files")

    template_count = len(list(templates_dir.glob("*.docx"))) if templates_dir.exists() else 0

    print(f"[OK] Output folder: {output_dir.absolute()}")
    print(f"[OK] Database folder: {database_dir.absolute()}")
    print(f"[OK] Templates folder: {templates_dir.absolute()} ({template_count} templates)")
    print(f"[OK] Config folder: {config_dir.absolute()}")

    if template_count == 0:
        print(f"\n[WARNING] WARNING: Word templates not found!")
        print(f"   Add .docx files to templates folder for document generation")


class PDFProcessorApp:
    def __init__(self, root):
        self.root = root

        # Load UI configuration from config file
        try:
            ui_config = config_loader.load('ui_settings')
            colors_from_config = ui_config.get('colors', {})

            # Map config color keys to internal keys (handle 'background' -> 'bg')
            self.colors = {
                'bg': colors_from_config.get('background', '#ECE9D8'),
                'primary': colors_from_config.get('primary', '#0054E3'),
                'secondary': colors_from_config.get('secondary', '#D4D0C8'),
                'accent': colors_from_config.get('accent', '#B0B0B0'),
                'text': colors_from_config.get('text', '#000000'),
                'border': colors_from_config.get('border', '#808080'),
                'button': colors_from_config.get('button', '#D4D0C8'),
                'button_active': colors_from_config.get('button_active', '#C0C0C0')
            }

            window_config = ui_config.get('window', {})
            self.button_labels = ui_config.get('button_labels', {})

            # Set window properties from config
            from src.version_info import get_short_version_string
            version_string = get_short_version_string()
            
            base_title = window_config.get('title', 'Парсер PDF и генератор Word документов - ОАТИ')
            window_title = f"{base_title} - {version_string}"
            window_size = f"{window_config.get('width', 1100)}x{window_config.get('height', 750)}"

            self.root.title(window_title)
            self.root.geometry(window_size)

        except Exception as e:
            print(f"[WARNING] Failed to load UI config: {e}")
            print(f"[WARNING] Using default settings")
            # Fallback to default values
            self.colors = {
                'bg': '#ECE9D8',
                'primary': '#0054E3',
                'secondary': '#D4D0C8',
                'accent': '#B0B0B0',
                'text': '#000000',
                'border': '#808080',
                'button': '#D4D0C8',
                'button_active': '#C0C0C0'
            }
            self.button_labels = {
                'select_pdf': 'Выбрать PDF',
                'process': 'Обработать',
                'clear': 'Очистить',
                'directors': 'Директора',
                'inspector_chiefs': 'Начальники инспекций'
            }
            
            from src.version_info import get_short_version_string
            version_string = get_short_version_string()
            self.root.title(f"Парсер PDF и генератор Word документов - ОАТИ - {version_string}")
            self.root.geometry("1100x750")

        self.root.configure(bg=self.colors['bg'])

        self.processing_queue = []
        self.document_outputs = []
        self.processing_runs = []
        self.listbox_index_map = []
        
        # Core services (always needed)
        self.parser = PDFParser()
        self.db = Database()
        self.decliner = NameDeclension(database=self.db)
        self.generator = WordGenerator()
        
        # Lazy-loaded services (initialized on demand)
        self.russia_post = None
        self.scanner = None
        self.mosedo_automation = None
        self.ai_assistant = None
        
        # MOSEDO workflow state
        self.current_workflow_id = None
        self.current_workflow_name = None
        self.current_steps = []
        
        # Track which services have been initialized
        self.services_initialized = {
            'russia_post': False,
            'scanner': False,
            'mosedo': False,
            'ai_assistant': False
        }

        self.setup_ui()
    
    def init_russia_post(self):
        """Lazy initialize Russia Post API service."""
        if not self.services_initialized['russia_post']:
            try:
                self.log("Инициализация сервиса Почты России...")
                from src.russia_post import RussiaPostAPI
                self.russia_post = RussiaPostAPI()
                self.services_initialized['russia_post'] = True
                self.log("✓ Сервис Почты России загружен")
            except Exception as e:
                self.log(f"✗ Ошибка при инициализации Почты России: {e}")
                self.log("  Модуль будет работать в offline режиме")
                self.log("  Повторная попытка будет выполнена при следующем обращении")
                # Keep russia_post as None - UI will check for None before using
                # DO NOT set services_initialized['russia_post'] = True here!
                # This allows retry on next attempt (e.g., after user fixes credentials)
                self.russia_post = None
    
    def init_scanner(self):
        """Lazy initialize Scanner service."""
        if not self.services_initialized['scanner']:
            self.log_scanner("Инициализация сервиса сканера...")
            from src.scanner_integration import ScannerService
            self.scanner = ScannerService()
            self.services_initialized['scanner'] = True
            self.log_scanner("✓ Сервис сканера загружен")
    
    def init_ai_assistant(self):
        """Lazy initialize AI Assistant service."""
        if not self.services_initialized['ai_assistant']:
            try:
                self.log("Инициализация AI Помощника...")
                from src.ai_assistant import AIAssistant
                self.ai_assistant = AIAssistant()
                self.services_initialized['ai_assistant'] = True
                status = self.ai_assistant.get_status()
                if status['is_online']:
                    self.log(f"✓ AI Помощник загружен и готов ({status['provider']})")
                else:
                    self.log(f"⚠ AI Помощник загружен в offline режиме: {status['offline_reason']}")
            except Exception as e:
                self.log(f"✗ Ошибка при инициализации AI Помощника: {e}")
                self.ai_assistant = None
    
    def init_mosedo(self):
        """Lazy initialize MOSEDO automation with workflow builder."""
        if not self.services_initialized['mosedo']:
            self.log_mosedo("Инициализация МОСЭДО автоматизации...")
            from src.mosedo_automation import MOSEDOAutomation
            self.mosedo_automation = MOSEDOAutomation(database=self.db)
            self.services_initialized['mosedo'] = True
            self.log_mosedo("✓ МОСЭДО автоматизация загружена")
            self.load_mosedo_credentials()
            self.load_workflows_list()

    def on_tab_changed(self, event):
        """Handle notebook tab change event for lazy loading."""
        selected_tab = event.widget.select()
        tab_index = event.widget.index(selected_tab)
        
        # Tab indices: 0=Уведомления, 1=ЭЗП, 2=Почта России, 3=Сканирование, 4=МОСЭДО
        if tab_index == 2 and not self.services_initialized['russia_post']:
            self.init_russia_post()
            self.load_shipments_table()  # Load demo data after init
        elif tab_index == 3 and not self.services_initialized['scanner']:
            self.init_scanner()
        elif tab_index == 4 and not self.services_initialized['mosedo']:
            self.init_mosedo()

    def setup_ui(self):
        main_frame = tk.Frame(self.root, bg=self.colors['bg'], padx=15, pady=15)
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)

        header_frame = tk.Frame(main_frame, bg=self.colors['primary'], height=60)
        header_frame.grid(row=0, column=0, pady=(0, 15), sticky=(tk.W, tk.E))
        header_frame.grid_propagate(False)

        title_label = tk.Label(header_frame, 
                              text="Парсер PDF и генератор Word документов", 
                              font=('Segoe UI', 18, 'bold'),
                              bg=self.colors['primary'],
                              fg='white')
        title_label.pack(pady=15)

        notebook = ttk.Notebook(main_frame)
        notebook.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        notifications_tab = tk.Frame(notebook, bg=self.colors['bg'])
        ezp_tab = tk.Frame(notebook, bg=self.colors['bg'])
        russia_post_tab = tk.Frame(notebook, bg=self.colors['bg'])
        scanner_tab = tk.Frame(notebook, bg=self.colors['bg'])
        mosedo_tab = tk.Frame(notebook, bg=self.colors['bg'])
        ai_tab = tk.Frame(notebook, bg=self.colors['bg'])

        notebook.add(notifications_tab, text='  Уведомления  ')
        notebook.add(ezp_tab, text='  Работа с ЭЗП  ')
        notebook.add(russia_post_tab, text='  Почта России  ')
        notebook.add(scanner_tab, text='  Сканирование  ')
        notebook.add(mosedo_tab, text='  МОСЭДО Робот  ')
        notebook.add(ai_tab, text='  🤖 AI Помощник  ')
        
        # Bind tab change event for lazy loading
        notebook.bind("<<NotebookTabChanged>>", self.on_tab_changed)

        self.setup_notifications_tab(notifications_tab)
        self.setup_ai_tab(ai_tab)
        self.setup_ezp_tab(ezp_tab)
        self.setup_russia_post_tab(russia_post_tab)
        self.setup_scanner_tab(scanner_tab)
        self.setup_mosedo_tab(mosedo_tab)

        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(main_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.grid(row=2, column=0, pady=(10, 5), sticky=(tk.W, tk.E))

        status_frame = tk.Frame(main_frame, bg='#34495e', height=35)
        status_frame.grid(row=3, column=0, sticky=(tk.W, tk.E))
        status_frame.grid_propagate(False)

        self.status_label = tk.Label(status_frame, 
                                     text="[STATUS] Готово к работе | Шаблонов: загружаются...",
                                     font=('Segoe UI', 9),
                                     bg='#34495e',
                                     fg='white',
                                     anchor=tk.W,
                                     padx=15)
        self.status_label.pack(fill=tk.BOTH, expand=True)

        self.update_status()

    def setup_notifications_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        parent.rowconfigure(2, weight=1)

        button_frame = tk.Frame(parent, bg=self.colors['bg'])
        button_frame.grid(row=0, column=0, pady=(15, 15), sticky=(tk.W, tk.E))

        btn_style = {
            'font': ('Tahoma', 9, 'bold'),
            'relief': tk.RAISED,
            'cursor': 'hand2',
            'padx': 20,
            'pady': 12,
            'bd': 2
        }

        self.btn_select = tk.Button(button_frame, text=self.button_labels.get('select_pdf', 'Выбрать PDF'), 
                                    command=self.select_files,
                                    bg=self.colors['button'], fg=self.colors['text'],
                                    activebackground=self.colors['button_active'], **btn_style)
        self.btn_select.grid(row=0, column=0, padx=5)

        self.btn_process = tk.Button(button_frame, text=self.button_labels.get('process', 'Обработать'), 
                                     command=self.process_files,
                                     bg=self.colors['button'], fg=self.colors['text'],
                                     activebackground=self.colors['button_active'], **btn_style)
        self.btn_process.grid(row=0, column=1, padx=5)

        self.btn_clear = tk.Button(button_frame, text=self.button_labels.get('clear', 'Очистить'), 
                                   command=self.clear_files,
                                   bg=self.colors['button'], fg=self.colors['text'],
                                   activebackground=self.colors['button_active'], **btn_style)
        self.btn_clear.grid(row=0, column=2, padx=5)
        
        self.btn_organizations = tk.Button(button_frame, text='Организации', 
                                           command=self.open_organizations_window,
                                           bg=self.colors['button'], fg=self.colors['text'],
                                           activebackground=self.colors['button_active'], **btn_style)
        self.btn_organizations.grid(row=1, column=0, padx=5, pady=(10, 0))

        self.btn_depts = tk.Button(button_frame, text=self.button_labels.get('directors', 'Директора'), 
                                   command=self.open_departments_window,
                                   bg=self.colors['button'], fg=self.colors['text'],
                                   activebackground=self.colors['button_active'], **btn_style)
        self.btn_depts.grid(row=0, column=3, padx=5)

        self.btn_inspectors = tk.Button(button_frame, text=self.button_labels.get('inspector_chiefs', 'Начальники инспекций'), 
                                        command=self.open_inspector_chiefs_window,
                                        bg=self.colors['button'], fg=self.colors['text'],
                                        activebackground=self.colors['button_active'], **btn_style)
        self.btn_inspectors.grid(row=0, column=4, padx=5)
        
        self.btn_history = tk.Button(button_frame, text='История обработки', 
                                      command=self.open_history_window,
                                      bg=self.colors['button'], fg=self.colors['text'],
                                      activebackground=self.colors['button_active'], **btn_style)
        self.btn_history.grid(row=0, column=5, padx=5)

        paned_window = tk.PanedWindow(parent, orient=tk.HORIZONTAL, 
                                      bg=self.colors['bg'], 
                                      sashwidth=8, 
                                      sashrelief=tk.RAISED)
        paned_window.grid(row=1, column=0, pady=(0, 10), sticky=(tk.W, tk.E, tk.N, tk.S))

        files_frame = tk.LabelFrame(paned_window, 
                                   text="  PDF файлы для обработки  ",
                                   font=('Segoe UI', 11, 'bold'),
                                   bg=self.colors['bg'],
                                   fg=self.colors['text'],
                                   relief=tk.GROOVE,
                                   bd=2)
        files_frame.columnconfigure(0, weight=1)
        files_frame.rowconfigure(0, weight=1)

        drop_zone = tk.Frame(files_frame, bg='#e8f4f8', relief=tk.SOLID, bd=2)
        drop_zone.grid(row=0, column=0, padx=10, pady=10, sticky=(tk.W, tk.E, tk.N, tk.S))
        drop_zone.columnconfigure(0, weight=1)
        drop_zone.rowconfigure(0, weight=1)

        scrollbar = tk.Scrollbar(drop_zone, orient=tk.VERTICAL)
        self.files_listbox = tk.Listbox(drop_zone, 
                                        yscrollcommand=scrollbar.set, 
                                        height=8,
                                        bg='white',
                                        font=('Segoe UI', 10),
                                        selectbackground=self.colors['primary'],
                                        relief=tk.FLAT,
                                        bd=0)
        scrollbar.config(command=self.files_listbox.yview)

        self.files_listbox.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=5)
        
        self.files_listbox.bind('<<ListboxSelect>>', self.on_pdf_select)

        drop_label = tk.Label(files_frame, 
                            text="Перетащите PDF файлы сюда или используйте кнопку 'Выбрать PDF'",
                            font=('Segoe UI', 9, 'italic'),
                            bg=self.colors['bg'],
                            fg='#7f8c8d')
        drop_label.grid(row=1, column=0, pady=(0, 10))

        preview_frame = tk.LabelFrame(paned_window, 
                                      text="  Предпросмотр документа  ",
                                      font=('Segoe UI', 11, 'bold'),
                                      bg=self.colors['bg'],
                                      fg=self.colors['text'],
                                      relief=tk.GROOVE,
                                      bd=2)
        preview_frame.columnconfigure(0, weight=1)
        preview_frame.rowconfigure(0, weight=1)

        preview_scrollbar = tk.Scrollbar(preview_frame, orient=tk.VERTICAL)
        self.preview_text = tk.Text(preview_frame, 
                                    yscrollcommand=preview_scrollbar.set,
                                    bg='white',
                                    fg=self.colors['text'],
                                    font=('Segoe UI', 10),
                                    relief=tk.FLAT,
                                    state=tk.DISABLED,
                                    padx=15,
                                    pady=15,
                                    wrap=tk.WORD)
        preview_scrollbar.config(command=self.preview_text.yview)

        self.preview_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        preview_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=10)

        paned_window.add(files_frame, minsize=300)
        paned_window.add(preview_frame, minsize=400)

        log_frame = tk.LabelFrame(parent, 
                                 text="  Журнал обработки  ",
                                 font=('Segoe UI', 11, 'bold'),
                                 bg=self.colors['bg'],
                                 fg=self.colors['text'],
                                 relief=tk.GROOVE,
                                 bd=2)
        log_frame.grid(row=2, column=0, pady=(0, 10), sticky=(tk.W, tk.E, tk.N, tk.S))
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)

        log_scrollbar = tk.Scrollbar(log_frame, orient=tk.VERTICAL)
        self.log_text = tk.Text(log_frame, 
                               height=8, 
                               yscrollcommand=log_scrollbar.set,
                               bg='#2c3e50',
                               fg='#ecf0f1',
                               font=('Consolas', 9),
                               relief=tk.FLAT,
                               state=tk.DISABLED,
                               padx=10,
                               pady=10)
        log_scrollbar.config(command=self.log_text.yview)

        self.log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        log_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=10)

        try:
            self.files_listbox.drop_target_register(tkdnd.DND_FILES)
            self.files_listbox.dnd_bind('<<Drop>>', self.on_drop)
        except:
            self.log("Drag-and-drop недоступен в этой среде")

    def setup_ai_tab(self, parent):
        """Setup AI Assistant tab with status indicator and settings."""
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)
        
        # Header
        header_frame = tk.Frame(parent, bg=self.colors['primary'], relief=tk.RAISED, bd=2)
        header_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        tk.Label(header_frame, text="🤖 YandexGPT AI Помощник", 
                font=('Segoe UI', 14, 'bold'), bg=self.colors['primary'], 
                fg='white', pady=10).pack()
        
        # Main content frame
        content_frame = tk.Frame(parent, bg=self.colors['bg'])
        content_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10)
        content_frame.columnconfigure(0, weight=1)
        
        # Status Section
        status_frame = tk.LabelFrame(content_frame, text="  Статус  ", font=('Segoe UI', 10, 'bold'),
                                    bg=self.colors['bg'], fg=self.colors['text'], relief=tk.GROOVE, bd=2)
        status_frame.grid(row=0, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        status_inner = tk.Frame(status_frame, bg=self.colors['bg'])
        status_inner.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(status_inner, text="Подключение:", font=('Tahoma', 9), 
                bg=self.colors['bg']).grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.ai_status_label = tk.Label(status_inner, text="🔴 Offline", font=('Tahoma', 9, 'bold'),
                                       bg=self.colors['bg'], fg='red')
        self.ai_status_label.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        tk.Button(status_inner, text="Проверить подключение", command=self.test_ai_connection,
                 font=('Tahoma', 9, 'bold'), bg='#3498db', fg='white', 
                 relief=tk.RAISED, bd=2, padx=15, pady=8).grid(row=0, column=2, padx=10)
        
        # Settings Section
        settings_frame = tk.LabelFrame(content_frame, text="  Настройки YandexGPT API  ", 
                                      font=('Segoe UI', 10, 'bold'),
                                      bg=self.colors['bg'], fg=self.colors['text'], 
                                      relief=tk.GROOVE, bd=2)
        settings_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(0, 10))
        
        settings_inner = tk.Frame(settings_frame, bg=self.colors['bg'])
        settings_inner.pack(fill=tk.X, padx=10, pady=10)
        
        tk.Label(settings_inner, text="API Key:", font=('Tahoma', 9), 
                bg=self.colors['bg']).grid(row=0, column=0, sticky=tk.W, pady=5)
        self.ai_api_key_entry = tk.Entry(settings_inner, font=('Tahoma', 9), show='*', width=40)
        self.ai_api_key_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=5, pady=5)
        tk.Button(settings_inner, text="📋 Вставить", command=self.paste_api_key,
                 font=('Tahoma', 8), bg='#95a5a6', fg='white', 
                 relief=tk.RAISED, bd=1, padx=8, pady=4).grid(row=0, column=2, padx=5)
        
        tk.Label(settings_inner, text="Folder ID:", font=('Tahoma', 9), 
                bg=self.colors['bg']).grid(row=1, column=0, sticky=tk.W, pady=5)
        self.ai_folder_id_entry = tk.Entry(settings_inner, font=('Tahoma', 9), width=40)
        self.ai_folder_id_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), padx=5, pady=5)
        tk.Button(settings_inner, text="📋 Вставить", command=self.paste_folder_id,
                 font=('Tahoma', 8), bg='#95a5a6', fg='white', 
                 relief=tk.RAISED, bd=1, padx=8, pady=4).grid(row=1, column=2, padx=5)
        
        settings_inner.columnconfigure(1, weight=1)
        
        btn_frame = tk.Frame(settings_inner, bg=self.colors['bg'])
        btn_frame.grid(row=2, column=0, columnspan=2, pady=10)
        
        tk.Button(btn_frame, text="💾 Сохранить credentials", command=self.save_ai_credentials,
                 font=('Tahoma', 9, 'bold'), bg='#27ae60', fg='white', 
                 relief=tk.RAISED, bd=2, padx=15, pady=8).pack(side=tk.LEFT, padx=5)
        tk.Button(btn_frame, text="🗑 Очистить", command=self.clear_ai_credentials,
                 font=('Tahoma', 9, 'bold'), bg='#e74c3c', fg='white', 
                 relief=tk.RAISED, bd=2, padx=15, pady=8).pack(side=tk.LEFT, padx=5)
        
        # Log Section  
        log_frame = tk.LabelFrame(content_frame, text="  Журнал  ", font=('Segoe UI', 10, 'bold'),
                                 bg=self.colors['bg'], fg=self.colors['text'], relief=tk.GROOVE, bd=2)
        log_frame.grid(row=2, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        log_frame.rowconfigure(0, weight=1)
        log_frame.columnconfigure(0, weight=1)
        content_frame.rowconfigure(2, weight=1)
        
        log_scrollbar = tk.Scrollbar(log_frame, orient=tk.VERTICAL)
        self.ai_log = tk.Text(log_frame, yscrollcommand=log_scrollbar.set, bg='#2c3e50', 
                             fg='#ecf0f1', font=('Consolas', 9), relief=tk.FLAT, 
                             state=tk.DISABLED, padx=10, pady=10, wrap=tk.WORD, height=15)
        log_scrollbar.config(command=self.ai_log.yview)
        
        self.ai_log.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
        log_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=5)
        
        # Initialize AI on tab creation
        self.init_ai_assistant()
        if self.ai_assistant:
            status = self.ai_assistant.get_status()
            if status['is_online']:
                self.ai_status_label.config(text="🟢 Online", fg='green')
                self.log_ai(f"✓ AI подключен: {status['model']}")
            else:
                self.log_ai(f"⚠ AI offline: {status['offline_reason']}")

    def setup_ezp_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)  # Log frame expands

        btn_style = {
            'font': ('Tahoma', 9, 'bold'),
            'relief': tk.RAISED,
            'cursor': 'hand2',
            'padx': 20,
            'pady': 12,
            'bd': 2
        }

        # PDF Drag & Drop zone
        pdf_frame = tk.LabelFrame(parent, 
                                 text="  PDF файлы для обработки  ",
                                 font=('Segoe UI', 10, 'bold'),
                                 bg=self.colors['bg'],
                                 fg=self.colors['text'],
                                 relief=tk.GROOVE,
                                 bd=2)
        pdf_frame.grid(row=0, column=0, pady=(15, 10), sticky=(tk.W, tk.E), padx=10)
        pdf_frame.columnconfigure(0, weight=1)
        
        # Drag & Drop zone
        self.ezp_drop_zone = tk.Label(pdf_frame,
                                     text="📄 Перетащите PDF файлы сюда\n"
                                          "или кликните для выбора файлов\n\n"
                                          "Поддерживается batch обработка",
                                     bg='#ecf0f1',
                                     fg='#7f8c8d',
                                     font=('Segoe UI', 11),
                                     relief=tk.SUNKEN,
                                     bd=2,
                                     height=5,
                                     cursor='hand2')
        self.ezp_drop_zone.grid(row=0, column=0, padx=10, pady=10, sticky=(tk.W, tk.E))
        
        # Initialize PDF files list
        self.ezp_pdf_files = []
        
        # Register drag & drop
        try:
            from tkinterdnd2 import DND_FILES
            self.ezp_drop_zone.drop_target_register(DND_FILES)
            self.ezp_drop_zone.dnd_bind('<<Drop>>', self.ezp_on_drop)
        except:
            pass  # tkinterdnd2 not available
        
        # Click to browse FILES (not folder)
        self.ezp_drop_zone.bind('<Button-1>', lambda e: self.load_pdf_files())
        
        # Process button & progress
        process_frame = tk.Frame(parent, bg=self.colors['bg'])
        process_frame.grid(row=1, column=0, pady=(0, 10), sticky=(tk.W, tk.E), padx=10)
        process_frame.columnconfigure(1, weight=1)

        self.btn_process_ezp = tk.Button(process_frame, text='🚀 Обработать ЭЗП (AI)', 
                                         command=self.process_ezp_batch,
                                         bg='#27ae60', fg='white',
                                         activebackground='#229954', 
                                         font=('Tahoma', 10, 'bold'),
                                         relief=tk.RAISED,
                                         cursor='hand2',
                                         padx=30,
                                         pady=15,
                                         bd=3)
        self.btn_process_ezp.grid(row=0, column=0, padx=(10, 10), sticky=tk.W)
        
        # Progress bar
        try:
            from tkinter import ttk
            self.ezp_progress = ttk.Progressbar(process_frame, 
                                               mode='determinate',
                                               length=300)
            self.ezp_progress.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=10)
            
            self.ezp_progress_label = tk.Label(process_frame, 
                                              text="Готов к обработке",
                                              bg=self.colors['bg'],
                                              fg=self.colors['text'],
                                              font=('Segoe UI', 9))
            self.ezp_progress_label.grid(row=0, column=2, padx=10, sticky=tk.E)
        except:
            self.ezp_progress = None
            self.ezp_progress_label = None

        ezp_log_frame = tk.LabelFrame(parent, 
                                     text="  Журнал обработки ЭЗП  ",
                                     font=('Segoe UI', 11, 'bold'),
                                     bg=self.colors['bg'],
                                     fg=self.colors['text'],
                                     relief=tk.GROOVE,
                                     bd=2)
        ezp_log_frame.grid(row=2, column=0, pady=(0, 10), sticky=(tk.W, tk.E, tk.N, tk.S), padx=10)
        ezp_log_frame.columnconfigure(0, weight=1)
        ezp_log_frame.rowconfigure(0, weight=1)

        ezp_log_scrollbar = tk.Scrollbar(ezp_log_frame, orient=tk.VERTICAL)
        self.ezp_log_text = tk.Text(ezp_log_frame, 
                                    height=12, 
                                    yscrollcommand=ezp_log_scrollbar.set,
                                    bg='#2c3e50',
                                    fg='#ecf0f1',
                                    font=('Consolas', 9),
                                    relief=tk.FLAT,
                                    state=tk.DISABLED,
                                    padx=10,
                                    pady=10)
        ezp_log_scrollbar.config(command=self.ezp_log_text.yview)

        self.ezp_log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=5, pady=5)
        ezp_log_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=5)

        self.ezp_pdf_folder = None
        
        self.ezp_log("✓ Модуль ЭЗП готов к работе")
        self.ezp_log("📋 Инструкция:")
        self.ezp_log("  1. Перетащите PDF файлы в зону выше")
        self.ezp_log("  2. Нажмите 'Обработать ЭЗП' - программа автоматически использует шаблон")
        self.ezp_log("  3. Получите готовый Excel файл в папке output/")
        self.ezp_log("")
        self.ezp_log(f"📁 Шаблон: templates/ezp_template.xls")

    def setup_russia_post_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(3, weight=1)  # Table frame gets expansion weight

        btn_style = {
            'font': ('Tahoma', 9, 'bold'),
            'relief': tk.RAISED,
            'cursor': 'hand2',
            'padx': 20,
            'pady': 12,
            'bd': 2
        }

        # Status indicator frame
        status_frame = tk.Frame(parent, bg=self.colors['bg'])
        status_frame.grid(row=0, column=0, pady=(15, 5), sticky=(tk.W, tk.E))
        
        tk.Label(status_frame, text="Статус API:", 
                font=('Tahoma', 9, 'bold'),
                bg=self.colors['bg'], fg=self.colors['text']).grid(row=0, column=0, padx=(10, 5))
        
        self.russia_post_status_label = tk.Label(status_frame, text="🟡 Не проверено", 
                                                font=('Tahoma', 9),
                                                bg=self.colors['bg'], fg=self.colors['text'])
        self.russia_post_status_label.grid(row=0, column=1, padx=5)

        # Credentials configuration frame
        cred_frame = tk.LabelFrame(parent, 
                                   text="  Настройки API  ",
                                   font=('Segoe UI', 10, 'bold'),
                                   bg=self.colors['bg'],
                                   fg=self.colors['text'],
                                   relief=tk.GROOVE,
                                   bd=2)
        cred_frame.grid(row=1, column=0, pady=(5, 10), sticky=(tk.W, tk.E), padx=10)
        
        # Login field
        tk.Label(cred_frame, text="Логин:", 
                font=('Tahoma', 9),
                bg=self.colors['bg'], fg=self.colors['text']).grid(row=0, column=0, padx=(10, 5), pady=10, sticky=tk.W)
        
        self.russia_post_login_entry = tk.Entry(cred_frame, font=('Tahoma', 9), width=30)
        self.russia_post_login_entry.grid(row=0, column=1, padx=5, pady=10, sticky=tk.W)
        
        # Password field
        tk.Label(cred_frame, text="Пароль:", 
                font=('Tahoma', 9),
                bg=self.colors['bg'], fg=self.colors['text']).grid(row=0, column=2, padx=(20, 5), pady=10, sticky=tk.W)
        
        self.russia_post_password_entry = tk.Entry(cred_frame, font=('Tahoma', 9), width=30, show='*')
        self.russia_post_password_entry.grid(row=0, column=3, padx=5, pady=10, sticky=tk.W)
        
        # Buttons for credentials management
        tk.Button(cred_frame, text='💾 Сохранить', 
                  command=self.save_russia_post_credentials,
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'],
                  font=('Tahoma', 9, 'bold'), relief=tk.RAISED, cursor='hand2',
                  padx=15, pady=8, bd=2).grid(row=0, column=4, padx=5, pady=10)
        
        tk.Button(cred_frame, text='🗑 Очистить', 
                  command=self.clear_russia_post_credentials,
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'],
                  font=('Tahoma', 9, 'bold'), relief=tk.RAISED, cursor='hand2',
                  padx=15, pady=8, bd=2).grid(row=0, column=5, padx=5, pady=10)

        button_frame = tk.Frame(parent, bg=self.colors['bg'])
        button_frame.grid(row=2, column=0, pady=(5, 15), sticky=(tk.W, tk.E))

        tk.Button(button_frame, text='🔍 Проверить подключение API', 
                  command=self.test_russia_post_connection,
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=0, padx=5)

        tk.Button(button_frame, text='Обновить статусы', 
                  command=self.refresh_all_shipments_statuses,
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=1, padx=5)

        tk.Button(button_frame, text='Добавить отправление', 
                  command=self.add_new_shipment_dialog,
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=2, padx=5)

        tk.Button(button_frame, text='Показать возвраты', 
                  command=self.show_returns_only,
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=3, padx=5)

        table_frame = tk.LabelFrame(parent, 
                                    text="  База отправленных писем  ",
                                    font=('Segoe UI', 11, 'bold'),
                                    bg=self.colors['bg'],
                                    fg=self.colors['text'],
                                    relief=tk.GROOVE,
                                    bd=2)
        table_frame.grid(row=3, column=0, pady=(0, 10), sticky=(tk.W, tk.E, tk.N, tk.S))
        table_frame.columnconfigure(0, weight=1)
        table_frame.rowconfigure(0, weight=1)

        scrollbar = tk.Scrollbar(table_frame, orient=tk.VERTICAL)
        
        self.russia_post_table = tk.Text(table_frame, 
                            height=15, 
                            yscrollcommand=scrollbar.set,
                            bg='white',
                            fg=self.colors['text'],
                            font=('Consolas', 9),
                            relief=tk.FLAT,
                            state=tk.NORMAL,
                            padx=10,
                            pady=10)
        scrollbar.config(command=self.russia_post_table.yview)

        self.russia_post_table.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=10)

        log_frame = tk.LabelFrame(parent, 
                                 text="  Журнал Почты России  ",
                                 font=('Segoe UI', 11, 'bold'),
                                 bg=self.colors['bg'],
                                 fg=self.colors['text'],
                                 relief=tk.GROOVE,
                                 bd=2)
        log_frame.grid(row=4, column=0, pady=(0, 10), sticky=(tk.W, tk.E, tk.N, tk.S))
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)

        log_scrollbar = tk.Scrollbar(log_frame, orient=tk.VERTICAL)
        self.russia_post_log = tk.Text(log_frame, 
                                       height=8, 
                                       yscrollcommand=log_scrollbar.set,
                                       bg='#2c3e50',
                                       fg='#ecf0f1',
                                       font=('Consolas', 9),
                                       relief=tk.FLAT,
                                       state=tk.DISABLED,
                                       padx=10,
                                       pady=10)
        log_scrollbar.config(command=self.russia_post_log.yview)

        self.russia_post_log.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        log_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=10)

        self.log_russia_post("✓ Вкладка Почты России готова")
        self.log_russia_post("Модуль загружается при первом использовании для ускорения запуска")
        
        # Try to load saved credentials
        self.load_russia_post_credentials()

    def setup_scanner_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1)

        btn_style = {
            'font': ('Tahoma', 9, 'bold'),
            'relief': tk.RAISED,
            'cursor': 'hand2',
            'padx': 20,
            'pady': 12,
            'bd': 2
        }

        # Status indicator frame
        status_frame = tk.Frame(parent, bg=self.colors['bg'])
        status_frame.grid(row=0, column=0, pady=(15, 5), sticky=(tk.W, tk.E))
        
        tk.Label(status_frame, text="Статус сканера:", 
                font=('Tahoma', 9, 'bold'),
                bg=self.colors['bg'], fg=self.colors['text']).grid(row=0, column=0, padx=(10, 5))
        
        self.scanner_status_label = tk.Label(status_frame, text="🟡 Не проверено", 
                                            font=('Tahoma', 9),
                                            bg=self.colors['bg'], fg=self.colors['text'])
        self.scanner_status_label.grid(row=0, column=1, padx=5)

        button_frame = tk.Frame(parent, bg=self.colors['bg'])
        button_frame.grid(row=1, column=0, pady=(5, 15), sticky=(tk.W, tk.E))

        tk.Button(button_frame, text='🔍 Проверить систему', 
                  command=self.check_scanner_system,
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=0, padx=5)

        tk.Button(button_frame, text='📥 Установить компоненты', 
                  command=self.launch_scanner_setup,
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=1, padx=5)

        tk.Button(button_frame, text='▶ Начать сканирование', 
                  command=lambda: self.log_scanner("Запуск сканера Canon DR-M260..."),
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=2, padx=5)

        tk.Button(button_frame, text='Настройки сканера', 
                  command=lambda: self.log_scanner("Открытие настроек сканера..."),
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=3, padx=5)

        tk.Button(button_frame, text='Обработать как обращение', 
                  command=lambda: self.log_scanner("Обработка скана через PDF парсер..."),
                  bg=self.colors['button'], fg=self.colors['text'],
                  activebackground=self.colors['button_active'], **btn_style).grid(row=0, column=4, padx=5)

        queue_frame = tk.LabelFrame(parent, 
                                   text="  Очередь отсканированных документов  ",
                                   font=('Segoe UI', 11, 'bold'),
                                   bg=self.colors['bg'],
                                   fg=self.colors['text'],
                                   relief=tk.GROOVE,
                                   bd=2)
        queue_frame.grid(row=2, column=0, pady=(0, 10), sticky=(tk.W, tk.E, tk.N, tk.S))
        queue_frame.columnconfigure(0, weight=1)
        queue_frame.rowconfigure(0, weight=1)

        scrollbar = tk.Scrollbar(queue_frame, orient=tk.VERTICAL)
        
        queue_text = tk.Text(queue_frame, 
                            height=10, 
                            yscrollcommand=scrollbar.set,
                            bg='white',
                            fg=self.colors['text'],
                            font=('Consolas', 9),
                            relief=tk.FLAT,
                            state=tk.NORMAL,
                            padx=10,
                            pady=10)
        scrollbar.config(command=queue_text.yview)

        queue_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=10)

        queue_text.insert('1.0', "ФАЙЛ                         СТРАНИЦ    OCR       ТИП ДОКУМЕНТА\n")
        queue_text.insert('end', "-" * 80 + "\n")
        queue_text.insert('end', "scan_16112025_143022.pdf     3          ✓         Обращение гражданина\n")
        queue_text.insert('end', "scan_16112025_144501.pdf     2          ✓         Ответ юриста\n")
        queue_text.config(state=tk.DISABLED)

        log_frame = tk.LabelFrame(parent, 
                                 text="  Журнал сканирования  ",
                                 font=('Segoe UI', 11, 'bold'),
                                 bg=self.colors['bg'],
                                 fg=self.colors['text'],
                                 relief=tk.GROOVE,
                                 bd=2)
        log_frame.grid(row=3, column=0, pady=(0, 10), sticky=(tk.W, tk.E, tk.N, tk.S))
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)

        log_scrollbar = tk.Scrollbar(log_frame, orient=tk.VERTICAL)
        self.scanner_log = tk.Text(log_frame, 
                                  height=8, 
                                  yscrollcommand=log_scrollbar.set,
                                  bg='#2c3e50',
                                  fg='#ecf0f1',
                                  font=('Consolas', 9),
                                  relief=tk.FLAT,
                                  state=tk.DISABLED,
                                  padx=10,
                                  pady=10)
        log_scrollbar.config(command=self.scanner_log.yview)

        self.scanner_log.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        log_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S), pady=10)

        self.log_scanner("✓ Модуль сканирования инициализирован")
        self.log_scanner("✓ Обнаружен сканер: Canon ImageFormula DR-M260")
        self.log_scanner("⚠ Драйвер TWAIN не подключен - требуется настройка")

    def setup_mosedo_tab(self, parent):
        """Setup MOSEDO Manual Workflow Builder with 3-panel layout."""
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(0, weight=1)
        
        main_paned = tk.PanedWindow(parent, orient=tk.HORIZONTAL, 
                                    sashwidth=5, bg=self.colors['bg'])
        main_paned.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        
        self.setup_workflows_panel(main_paned)
        self.setup_steps_editor_panel(main_paned)
        self.setup_controls_panel(main_paned)
        
        self.log_mosedo("✓ МОСЭДО Manual Workflow Builder загружен")
        self.log_mosedo("ℹ Создайте новый workflow или выберите существующий для редактирования")
    
    def setup_workflows_panel(self, parent):
        """Left panel: Workflows list."""
        left_frame = tk.Frame(parent, bg=self.colors['bg'], width=250)
        left_frame.pack_propagate(False)
        
        header = tk.Label(left_frame, text="📋 Workflows", 
                         font=('Segoe UI', 11, 'bold'),
                         bg=self.colors['bg'], fg=self.colors['text'])
        header.pack(pady=(10, 5))
        
        list_frame = tk.Frame(left_frame, bg=self.colors['bg'])
        list_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        scrollbar = tk.Scrollbar(list_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        self.workflows_listbox = tk.Listbox(list_frame, 
                                           yscrollcommand=scrollbar.set,
                                           font=('Segoe UI', 9),
                                           bg='white',
                                           selectmode=tk.SINGLE)
        self.workflows_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        self.workflows_listbox.bind('<<ListboxSelect>>', self.on_workflow_select)
        scrollbar.config(command=self.workflows_listbox.yview)
        
        btn_frame = tk.Frame(left_frame, bg=self.colors['bg'])
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        btn_style = {'font': ('Tahoma', 9), 'relief': tk.RAISED, 'bd': 2, 'pady': 5}
        
        tk.Button(btn_frame, text='➕ Создать', command=self.create_new_workflow,
                 bg='#27ae60', fg='white', **btn_style).pack(fill=tk.X, pady=2)
        tk.Button(btn_frame, text='✏ Переименовать', command=self.rename_workflow,
                 bg=self.colors['button'], fg=self.colors['text'], **btn_style).pack(fill=tk.X, pady=2)
        tk.Button(btn_frame, text='🗑 Удалить', command=self.delete_workflow,
                 bg='#e74c3c', fg='white', **btn_style).pack(fill=tk.X, pady=2)
        
        parent.add(left_frame)
    
    def setup_steps_editor_panel(self, parent):
        """Center panel: Steps editor with Treeview."""
        center_frame = tk.Frame(parent, bg=self.colors['bg'])
        
        header = tk.Label(center_frame, text="🔧 Редактор шагов workflow", 
                         font=('Segoe UI', 11, 'bold'),
                         bg=self.colors['bg'], fg=self.colors['text'])
        header.pack(pady=(10, 5))
        
        tree_frame = tk.Frame(center_frame, bg=self.colors['bg'])
        tree_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        scrollbar = tk.Scrollbar(tree_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        columns = ('num', 'action', 'selector', 'value', 'description')
        self.steps_tree = ttk.Treeview(tree_frame, columns=columns, 
                                      show='headings', yscrollcommand=scrollbar.set)
        
        self.steps_tree.heading('num', text='№')
        self.steps_tree.heading('action', text='Action')
        self.steps_tree.heading('selector', text='Selector (CSS)')
        self.steps_tree.heading('value', text='Value')
        self.steps_tree.heading('description', text='Description')
        
        self.steps_tree.column('num', width=40, stretch=False)
        self.steps_tree.column('action', width=80, stretch=False)
        self.steps_tree.column('selector', width=150)
        self.steps_tree.column('value', width=150)
        self.steps_tree.column('description', width=200)
        
        self.steps_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.config(command=self.steps_tree.yview)
        
        controls_frame = tk.Frame(center_frame, bg=self.colors['bg'])
        controls_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(controls_frame, text="Action:", bg=self.colors['bg'], 
                font=('Tahoma', 9)).grid(row=0, column=0, padx=5, pady=2, sticky=tk.W)
        
        self.action_var = tk.StringVar(value='click')
        action_combo = ttk.Combobox(controls_frame, textvariable=self.action_var,
                                    values=['navigate', 'click', 'type', 'wait'],
                                    state='readonly', width=12)
        action_combo.grid(row=0, column=1, padx=5, pady=2)
        
        tk.Label(controls_frame, text="Selector:", bg=self.colors['bg'],
                font=('Tahoma', 9)).grid(row=0, column=2, padx=5, pady=2, sticky=tk.W)
        self.selector_entry = tk.Entry(controls_frame, width=20)
        self.selector_entry.grid(row=0, column=3, padx=5, pady=2)
        
        tk.Label(controls_frame, text="Value:", bg=self.colors['bg'],
                font=('Tahoma', 9)).grid(row=0, column=4, padx=5, pady=2, sticky=tk.W)
        self.value_entry = tk.Entry(controls_frame, width=20)
        self.value_entry.grid(row=0, column=5, padx=5, pady=2)
        
        tk.Label(controls_frame, text="Description:", bg=self.colors['bg'],
                font=('Tahoma', 9)).grid(row=1, column=0, padx=5, pady=2, sticky=tk.W)
        self.description_entry = tk.Entry(controls_frame, width=60)
        self.description_entry.grid(row=1, column=1, columnspan=5, padx=5, pady=2, sticky=(tk.W, tk.E))
        
        btn_frame = tk.Frame(center_frame, bg=self.colors['bg'])
        btn_frame.pack(fill=tk.X, padx=10, pady=5)
        
        btn_style = {'font': ('Tahoma', 9), 'relief': tk.RAISED, 'bd': 2}
        
        tk.Button(btn_frame, text='⬆ Вверх', command=self.move_step_up,
                 bg=self.colors['button'], fg=self.colors['text'], **btn_style).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_frame, text='⬇ Вниз', command=self.move_step_down,
                 bg=self.colors['button'], fg=self.colors['text'], **btn_style).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_frame, text='➕ Добавить шаг', command=self.add_step,
                 bg='#27ae60', fg='white', **btn_style).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_frame, text='✏ Изменить', command=self.edit_step,
                 bg='#3498db', fg='white', **btn_style).pack(side=tk.LEFT, padx=2)
        tk.Button(btn_frame, text='🗑 Удалить шаг', command=self.delete_step,
                 bg='#e74c3c', fg='white', **btn_style).pack(side=tk.LEFT, padx=2)
        
        parent.add(center_frame)
    
    def setup_controls_panel(self, parent):
        """Right panel: Preview, controls, and log."""
        right_frame = tk.Frame(parent, bg=self.colors['bg'], width=300)
        right_frame.pack_propagate(False)
        
        header = tk.Label(right_frame, text="⚙ Управление", 
                         font=('Segoe UI', 11, 'bold'),
                         bg=self.colors['bg'], fg=self.colors['text'])
        header.pack(pady=(10, 5))
        
        info_frame = tk.LabelFrame(right_frame, text="  Workflow Info  ",
                                  font=('Segoe UI', 9, 'bold'),
                                  bg=self.colors['bg'], fg=self.colors['text'],
                                  relief=tk.GROOVE, bd=2)
        info_frame.pack(fill=tk.X, padx=10, pady=5)
        
        self.workflow_name_label = tk.Label(info_frame, text="Название: ---",
                                           bg=self.colors['bg'], fg=self.colors['text'],
                                           font=('Segoe UI', 9), anchor=tk.W)
        self.workflow_name_label.pack(fill=tk.X, padx=10, pady=2)
        
        self.workflow_steps_label = tk.Label(info_frame, text="Шагов: 0",
                                            bg=self.colors['bg'], fg=self.colors['text'],
                                            font=('Segoe UI', 9), anchor=tk.W)
        self.workflow_steps_label.pack(fill=tk.X, padx=10, pady=2)
        
        self.workflow_status_label = tk.Label(info_frame, text="Статус: Не выбран",
                                             bg=self.colors['bg'], fg=self.colors['text'],
                                             font=('Segoe UI', 9), anchor=tk.W)
        self.workflow_status_label.pack(fill=tk.X, padx=10, pady=2)
        
        cred_frame = tk.LabelFrame(right_frame, text="  Учетные данные MOSEDO  ",
                                   font=('Segoe UI', 9, 'bold'),
                                   bg=self.colors['bg'], fg=self.colors['text'],
                                   relief=tk.GROOVE, bd=2)
        cred_frame.pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(cred_frame, text="Логин:", font=('Tahoma', 9),
                bg=self.colors['bg'], fg=self.colors['text']).grid(row=0, column=0, padx=5, pady=5, sticky=tk.W)
        self.mosedo_login_entry = tk.Entry(cred_frame, font=('Tahoma', 9), width=15)
        self.mosedo_login_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        
        tk.Label(cred_frame, text="Пароль:", font=('Tahoma', 9),
                bg=self.colors['bg'], fg=self.colors['text']).grid(row=1, column=0, padx=5, pady=5, sticky=tk.W)
        self.mosedo_password_entry = tk.Entry(cred_frame, font=('Tahoma', 9), width=15, show='*')
        self.mosedo_password_entry.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)
        
        cred_btn_frame = tk.Frame(cred_frame, bg=self.colors['bg'])
        cred_btn_frame.grid(row=2, column=0, columnspan=2, pady=5)
        
        tk.Button(cred_btn_frame, text='💾 Сохранить', command=self.save_mosedo_credentials,
                 bg=self.colors['button'], fg=self.colors['text'],
                 activebackground=self.colors['button_active'],
                 font=('Tahoma', 8, 'bold'), relief=tk.RAISED, padx=10, pady=5, bd=2).pack(side=tk.LEFT, padx=2)
        tk.Button(cred_btn_frame, text='🗑 Очистить', command=self.clear_mosedo_credentials,
                 bg=self.colors['button'], fg=self.colors['text'],
                 activebackground=self.colors['button_active'],
                 font=('Tahoma', 8, 'bold'), relief=tk.RAISED, padx=10, pady=5, bd=2).pack(side=tk.LEFT, padx=2)
        
        # Browser Recording Frame (New Feature!)
        record_frame = tk.LabelFrame(right_frame, text="  🎬 Автоматическая запись  ",
                                     font=('Segoe UI', 9, 'bold'),
                                     bg=self.colors['bg'], fg=self.colors['text'],
                                     relief=tk.GROOVE, bd=2)
        record_frame.pack(fill=tk.X, padx=10, pady=5)
        
        record_btn_style = {'font': ('Tahoma', 9, 'bold'), 'relief': tk.RAISED, 
                           'bd': 2, 'pady': 8}
        
        tk.Button(record_frame, text='🔴 Записать мой workflow', command=self.start_recording_workflow,
                 bg='#e74c3c', fg='white', **record_btn_style).pack(fill=tk.X, padx=10, pady=5)
        tk.Button(record_frame, text='⏹ Остановить запись', command=self.stop_recording_workflow,
                 bg='#95a5a6', fg='white', **record_btn_style).pack(fill=tk.X, padx=10, pady=5)
        
        tk.Label(record_frame, text="(Откроет браузер для записи действий)",
                font=('Tahoma', 8, 'italic'), bg=self.colors['bg'], 
                fg='#7f8c8d').pack(padx=10, pady=(0, 5))
        
        btn_frame = tk.Frame(right_frame, bg=self.colors['bg'])
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        btn_style = {'font': ('Tahoma', 10, 'bold'), 'relief': tk.RAISED, 
                    'bd': 2, 'pady': 10}
        
        tk.Button(btn_frame, text='💾 Сохранить', command=self.save_current_workflow,
                 bg='#3498db', fg='white', **btn_style).pack(fill=tk.X, pady=2)
        tk.Button(btn_frame, text='▶ Запустить автоматизацию', command=self.run_workflow,
                 bg='#27ae60', fg='white', **btn_style).pack(fill=tk.X, pady=2)
        tk.Button(btn_frame, text='⏹ Остановить автоматизацию', command=self.stop_workflow,
                 bg='#e74c3c', fg='white', **btn_style).pack(fill=tk.X, pady=2)
        
        log_frame = tk.LabelFrame(right_frame, text="  Журнал  ",
                                 font=('Segoe UI', 9, 'bold'),
                                 bg=self.colors['bg'], fg=self.colors['text'],
                                 relief=tk.GROOVE, bd=2)
        log_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)
        
        log_scrollbar = tk.Scrollbar(log_frame, orient=tk.VERTICAL)
        self.mosedo_log = tk.Text(log_frame, 
                                 yscrollcommand=log_scrollbar.set,
                                 bg='#2c3e50', fg='#ecf0f1',
                                 font=('Consolas', 9),
                                 relief=tk.FLAT,
                                 state=tk.DISABLED,
                                 padx=10, pady=10,
                                 wrap=tk.WORD)
        log_scrollbar.config(command=self.mosedo_log.yview)
        
        self.mosedo_log.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        log_scrollbar.pack(side=tk.RIGHT, fill=tk.Y, pady=5)
        
        parent.add(right_frame)

    def log_russia_post(self, message):
        if hasattr(self, 'russia_post_log'):
            self.russia_post_log.config(state=tk.NORMAL)
            self.russia_post_log.insert(tk.END, f"[{self.get_timestamp()}] {message}\n")
            self.russia_post_log.see(tk.END)
            self.russia_post_log.config(state=tk.DISABLED)
    
    def log_ai(self, message):
        """Log message to AI Assistant tab."""
        if hasattr(self, 'ai_log'):
            self.ai_log.config(state=tk.NORMAL)
            self.ai_log.insert(tk.END, f"[{self.get_timestamp()}] {message}\n")
            self.ai_log.see(tk.END)
            self.ai_log.config(state=tk.DISABLED)
    
    def test_ai_connection(self):
        """Test AI connection and update status indicator."""
        self.log_ai("Проверка подключения к YandexGPT API...")
        
        if not self.ai_assistant:
            self.init_ai_assistant()
        
        if self.ai_assistant:
            status = self.ai_assistant.get_status()
            if status['is_online']:
                self.ai_status_label.config(text="🟢 Online", fg='green')
                self.log_ai(f"✓ Подключение успешно: {status['model']}")
                messagebox.showinfo("Успешно", f"AI подключен и готов к работе.\n\nМодель: {status['model']}\nПровайдер: {status['provider']}", parent=self.root)
            else:
                self.ai_status_label.config(text="🔴 Offline", fg='red')
                self.log_ai(f"✗ Подключение не удалось: {status['offline_reason']}")
                messagebox.showwarning("Offline", f"AI недоступен:\n{status['offline_reason']}", parent=self.root)
        else:
            self.ai_status_label.config(text="🔴 Offline", fg='red')
            self.log_ai("✗ AI Assistant не инициализирован")
    
    def paste_api_key(self):
        """Paste API Key from clipboard into entry field."""
        try:
            clipboard_text = self.root.clipboard_get()
            self.ai_api_key_entry.delete(0, tk.END)
            self.ai_api_key_entry.insert(0, clipboard_text.strip())
            self.log_ai(f"✓ API Key вставлен из буфера обмена ({len(clipboard_text)} символов)")
        except tk.TclError:
            self.log_ai("✗ Буфер обмена пуст или недоступен")
            messagebox.showwarning("Ошибка", "Буфер обмена пуст.\n\nСкопируйте API Key и попробуйте снова.", parent=self.root)
    
    def paste_folder_id(self):
        """Paste Folder ID from clipboard into entry field."""
        try:
            clipboard_text = self.root.clipboard_get()
            self.ai_folder_id_entry.delete(0, tk.END)
            self.ai_folder_id_entry.insert(0, clipboard_text.strip())
            self.log_ai(f"✓ Folder ID вставлен из буфера обмена: {clipboard_text.strip()[:15]}***")
        except tk.TclError:
            self.log_ai("✗ Буфер обмена пуст или недоступен")
            messagebox.showwarning("Ошибка", "Буфер обмена пуст.\n\nСкопируйте Folder ID и попробуйте снова.", parent=self.root)
    
    def save_ai_credentials(self):
        """Save YandexGPT credentials to secure storage."""
        from src.credential_manager import CredentialManager
        
        api_key = self.ai_api_key_entry.get().strip()
        folder_id = self.ai_folder_id_entry.get().strip()
        
        if not api_key or not folder_id:
            messagebox.showwarning("Недостаточно данных", 
                                  "Введите API Key и Folder ID для сохранения.", 
                                  parent=self.root)
            return
        
        success = CredentialManager.save_yandexgpt_credentials(api_key, folder_id)
        if success:
            self.log_ai(f"✓ Credentials сохранены в Windows Credential Manager (folder: {folder_id[:8]}***)")
            messagebox.showinfo("Успешно!", 
                              "✅ Credentials успешно сохранены\n\n"
                              "Ключи API сохранены в защищенном\n"
                              "Windows Credential Manager.\n\n"
                              "Теперь можно закрыть это окно.",
                              parent=self.root)
            # Reinitialize AI with new credentials
            self.services_initialized['ai_assistant'] = False
            self.ai_assistant = None
            self.init_ai_assistant()
            # Update status
            if self.ai_assistant:
                status = self.ai_assistant.get_status()
                if status['is_online']:
                    self.ai_status_label.config(text="🟢 Online", fg='green')
                    self.log_ai(f"✓ AI переподключен: {status['model']}")
        else:
            self.log_ai("✗ Ошибка доступа к Windows Credential Manager")
            result = messagebox.showerror("Ошибка сохранения", 
                               "❌ Windows Credential Manager недоступен\n\n"
                               "📋 ЧТО ДЕЛАТЬ:\n\n"
                               "1️⃣ Закройте программу\n"
                               "2️⃣ Правый клик на main.py или .exe\n"
                               "3️⃣ Выберите 'Запуск от имени администратора'\n"
                               "4️⃣ Попробуйте сохранить credentials снова\n\n"
                               "⚠️ ВАЖНО: Без прав администратора Windows\n"
                               "не позволяет сохранять ключи в безопасное\n"
                               "хранилище Credential Manager.\n\n"
                               "💡 АЛЬТЕРНАТИВА: Вставьте ключи каждый раз\n"
                               "при запуске (используйте кнопки 'Вставить').",
                               parent=self.root)
    
    def clear_ai_credentials(self):
        """Clear YandexGPT credentials from UI and secure storage."""
        from src.credential_manager import CredentialManager
        
        # Clear UI fields
        self.ai_api_key_entry.delete(0, tk.END)
        self.ai_folder_id_entry.delete(0, tk.END)
        
        # Ask for confirmation to delete from keyring
        confirm = messagebox.askyesno("Подтверждение", 
                                     "Удалить сохраненные credentials YandexGPT из защищенного хранилища?", 
                                     parent=self.root)
        
        if confirm:
            success = CredentialManager.delete_yandexgpt_credentials()
            if success:
                self.log_ai("✓ Credentials удалены из защищенного хранилища")
                self.ai_status_label.config(text="🔴 Offline", fg='red')
                # Reset AI
                self.services_initialized['ai_assistant'] = False
                self.ai_assistant = None
            else:
                self.log_ai("⚠ Credentials не найдены или уже удалены")
        else:
            self.log_ai("✓ Поля очищены (хранилище не изменено)")

    def save_russia_post_credentials(self):
        """Save Russia Post API credentials to secure storage."""
        login = self.russia_post_login_entry.get().strip()
        password = self.russia_post_password_entry.get().strip()
        
        if not login or not password:
            messagebox.showwarning(
                "Недостаточно данных",
                "Введите логин и пароль для сохранения.",
                parent=self.root
            )
            return
        
        success = CredentialManager.save_russia_post_credentials(login, password)
        if success:
            self.log_russia_post(f"✓ Credentials сохранены для пользователя: {login[:5]}***")
            messagebox.showinfo(
                "Успешно",
                f"Учетные данные для пользователя {login} успешно сохранены в защищенное хранилище Windows.",
                parent=self.root
            )
            # Reinitialize Russia Post API with new credentials
            # Force reinitialization by resetting the flag
            self.services_initialized['russia_post'] = False
            self.russia_post = None
            self.init_russia_post()
        else:
            self.log_russia_post("✗ Ошибка сохранения credentials")
            messagebox.showerror(
                "Ошибка",
                "Не удалось сохранить учетные данные. Проверьте доступ к Windows Credential Manager.",
                parent=self.root
            )
    
    def clear_russia_post_credentials(self):
        """Clear Russia Post API credentials from UI and secure storage."""
        login = self.russia_post_login_entry.get().strip()
        
        if not login:
            # Just clear UI fields
            self.russia_post_login_entry.delete(0, tk.END)
            self.russia_post_password_entry.delete(0, tk.END)
            self.log_russia_post("✓ Поля очищены")
            return
        
        # Ask for confirmation
        confirm = messagebox.askyesno(
            "Подтверждение",
            f"Удалить сохраненные учетные данные для пользователя {login}?",
            parent=self.root
        )
        
        if confirm:
            success = CredentialManager.delete_russia_post_credentials(login)
            self.russia_post_login_entry.delete(0, tk.END)
            self.russia_post_password_entry.delete(0, tk.END)
            
            if success:
                self.log_russia_post(f"✓ Credentials удалены для пользователя: {login}")
                messagebox.showinfo(
                    "Успешно",
                    f"Учетные данные для пользователя {login} удалены.",
                    parent=self.root
                )
            else:
                self.log_russia_post("⚠ Credentials не найдены или уже удалены")
    
    def load_russia_post_credentials(self):
        """Try to load Russia Post credentials from secure storage."""
        # Try to load from environment first (backwards compatibility)
        login = os.environ.get('RUSSIA_POST_LOGIN', '')
        
        if login:
            # Try to get password from keyring
            password = CredentialManager.get_russia_post_credentials(login)
            
            if password:
                self.russia_post_login_entry.delete(0, tk.END)
                self.russia_post_login_entry.insert(0, login)
                self.russia_post_password_entry.delete(0, tk.END)
                self.russia_post_password_entry.insert(0, password)
                self.log_russia_post(f"✓ Credentials загружены для пользователя: {login[:5]}***")
                return True
        
        return False

    def save_mosedo_credentials(self):
        """Save MOSEDO credentials to secure storage."""
        login = self.mosedo_login_entry.get().strip()
        password = self.mosedo_password_entry.get().strip()
        
        if not login or not password:
            messagebox.showwarning(
                "Недостаточно данных",
                "Введите логин и пароль для сохранения.",
                parent=self.root
            )
            return
        
        success = CredentialManager.save_mosedo_credentials(login, password)
        if success:
            self.log_mosedo(f"✓ Credentials сохранены для пользователя: {login[:5]}***")
            messagebox.showinfo(
                "Успешно",
                f"Учетные данные для пользователя {login} успешно сохранены в защищенное хранилище Windows.",
                parent=self.root
            )
        else:
            self.log_mosedo("✗ Ошибка сохранения credentials")
            messagebox.showerror(
                "Ошибка",
                "Не удалось сохранить учетные данные. Проверьте доступ к Windows Credential Manager.",
                parent=self.root
            )
    
    def clear_mosedo_credentials(self):
        """Clear MOSEDO credentials from UI and secure storage."""
        login = self.mosedo_login_entry.get().strip()
        
        if not login:
            self.mosedo_login_entry.delete(0, tk.END)
            self.mosedo_password_entry.delete(0, tk.END)
            self.log_mosedo("✓ Поля очищены")
            return
        
        confirm = messagebox.askyesno(
            "Подтверждение",
            f"Удалить сохраненные учетные данные для пользователя {login}?",
            parent=self.root
        )
        
        if confirm:
            success = CredentialManager.delete_mosedo_credentials(login)
            self.mosedo_login_entry.delete(0, tk.END)
            self.mosedo_password_entry.delete(0, tk.END)
            
            if success:
                self.log_mosedo(f"✓ Credentials удалены для пользователя: {login}")
                messagebox.showinfo(
                    "Успешно",
                    f"Учетные данные для пользователя {login} удалены.",
                    parent=self.root
                )
            else:
                self.log_mosedo("⚠ Credentials не найдены или уже удалены")
    
    def load_mosedo_credentials(self):
        """Try to load MOSEDO credentials from secure storage."""
        login = os.environ.get('MOSEDO_LOGIN', '')
        
        if not login:
            login = CredentialManager.get_saved_mosedo_login()
        
        if login:
            password = CredentialManager.get_mosedo_credentials(login)
            
            if password:
                self.mosedo_login_entry.delete(0, tk.END)
                self.mosedo_login_entry.insert(0, login)
                self.mosedo_password_entry.delete(0, tk.END)
                self.mosedo_password_entry.insert(0, password)
                self.log_mosedo(f"✓ Credentials загружены для пользователя: {login[:5]}***")
                return True
        
        return False

    def test_russia_post_connection(self):
        """Test Russia Post API connection and update status indicator."""
        def _test_thread():
            try:
                self.root.after(0, lambda: self.log_russia_post("🔍 Тестирование подключения к API Почты России..."))
                self.root.after(0, lambda: self.russia_post_status_label.config(text="🟡 Тестирование..."))
                
                # Ensure Russia Post service is initialized (synchronously in thread)
                if not self.services_initialized['russia_post']:
                    self.init_russia_post()
                
                # Check if russia_post is available
                if not self.russia_post:
                    self.root.after(0, lambda: self.russia_post_status_label.config(text="🔴 Ошибка"))
                    self.root.after(0, lambda: self.log_russia_post("✗ Модуль Russia Post не инициализирован"))
                    self.root.after(0, lambda: self.log_russia_post("  Возможная причина: ошибка при загрузке библиотек"))
                    return
                
                # Run connection test
                test_result = self.russia_post.test_connection()
                
                if test_result.get('connected'):
                    msg = test_result.get('message', 'Подключено')
                    details = test_result.get('details', '')
                    wsdl = test_result.get('wsdl_url', '')
                    login = test_result.get('login', '')
                    
                    self.root.after(0, lambda: self.russia_post_status_label.config(text="🟢 Подключено"))
                    self.root.after(0, lambda: self.log_russia_post(f"✓ {msg}"))
                    self.root.after(0, lambda: self.log_russia_post(f"  {details}"))
                    if wsdl:
                        self.root.after(0, lambda: self.log_russia_post(f"  WSDL: {wsdl}"))
                    if login:
                        self.root.after(0, lambda: self.log_russia_post(f"  Логин: {login}"))
                    if 'note' in test_result:
                        self.root.after(0, lambda: self.log_russia_post(f"  Примечание: {test_result['note']}"))
                else:
                    msg = test_result.get('message', 'Ошибка подключения')
                    details = test_result.get('details', '')
                    
                    self.root.after(0, lambda: self.russia_post_status_label.config(text="🔴 Ошибка"))
                    self.root.after(0, lambda: self.log_russia_post(f"✗ {msg}"))
                    self.root.after(0, lambda: self.log_russia_post(f"  {details}"))
                    
            except Exception as e:
                error_msg = str(e)
                self.root.after(0, lambda: self.russia_post_status_label.config(text="🔴 Ошибка"))
                self.root.after(0, lambda: self.log_russia_post(f"✗ Ошибка тестирования: {error_msg}"))
        
        threading.Thread(target=_test_thread, daemon=True).start()

    def refresh_all_shipments_statuses(self):
        def _refresh_thread():
            try:
                # Ensure Russia Post service is initialized (synchronously in thread)
                if not self.services_initialized['russia_post']:
                    self.init_russia_post()
                
                # Check if russia_post is available
                if not self.russia_post:
                    self.root.after(0, lambda: self.log_russia_post("✗ Модуль Russia Post не инициализирован"))
                    return
                
                self.root.after(0, lambda: self.log_russia_post("Начало обновления статусов отправлений..."))
                
                shipments = self.db.get_all_shipments()
                if not shipments:
                    self.root.after(0, lambda: self.log_russia_post("⚠ Нет отправлений в базе данных"))
                    return
                
                self.root.after(0, lambda: self.log_russia_post(f"Найдено отправлений: {len(shipments)}"))
                
                tracking_numbers = [s['tracking_number'] for s in shipments]
                self.root.after(0, lambda: self.log_russia_post(f"Запрос статусов через API Почты России..."))
                
                tracking_results = self.russia_post.batch_track(tracking_numbers)
                
                updated_count = 0
                for result in tracking_results:
                    tracking_number = result['tracking_number']
                    new_status = result['status']
                    
                    success = self.db.update_shipment_status(tracking_number, new_status)
                    if success:
                        updated_count += 1
                        self.root.after(0, lambda tn=tracking_number, st=new_status: self.log_russia_post(f"✓ {tn}: {st}"))
                    else:
                        self.root.after(0, lambda tn=tracking_number: self.log_russia_post(f"✗ Ошибка обновления {tn}"))
                
                self.root.after(0, lambda c=updated_count, t=len(shipments): self.log_russia_post(f"✓ Обновлено статусов: {c}/{t}"))
                
                self.root.after(0, self.load_shipments_table)
                
            except Exception as e:
                error_msg = str(e)
                self.root.after(0, lambda: self.log_russia_post(f"✗ Ошибка при обновлении: {error_msg}"))
        
        threading.Thread(target=_refresh_thread, daemon=True).start()

    def load_shipments_table(self, status_filter=None):
        try:
            shipments = self.db.get_all_shipments(status_filter=status_filter)
            
            self.russia_post_table.config(state=tk.NORMAL)
            self.russia_post_table.delete('1.0', tk.END)
            
            self.russia_post_table.insert('1.0', "ШПИ                 ПОЛУЧАТЕЛЬ              СТАТУС              ДАТА ОТПРАВКИ\n")
            self.russia_post_table.insert('end', "-" * 100 + "\n")
            
            if not shipments:
                self.russia_post_table.insert('end', "\n(Нет отправлений в базе данных)\n")
            else:
                for shipment in shipments:
                    tracking = shipment.get('tracking_number', 'N/A').ljust(20)
                    recipient = (shipment.get('recipient_name', 'N/A')[:20]).ljust(24)
                    status = (shipment.get('status', 'N/A')[:16]).ljust(20)
                    sent_date = shipment.get('sent_date', 'N/A').ljust(15)
                    
                    self.russia_post_table.insert('end', f"{tracking}{recipient}{status}{sent_date}\n")
            
            self.russia_post_table.config(state=tk.DISABLED)
            
        except Exception as e:
            self.log_russia_post(f"✗ Ошибка при загрузке таблицы: {str(e)}")

    def show_returns_only(self):
        try:
            self.log_russia_post("Фильтрация: показываем только возвраты...")
            
            all_shipments = self.db.get_all_shipments()
            returns = [s for s in all_shipments if 'Возврат' in s.get('status', '')]
            
            self.russia_post_table.config(state=tk.NORMAL)
            self.russia_post_table.delete('1.0', tk.END)
            
            self.russia_post_table.insert('1.0', "ШПИ                 ПОЛУЧАТЕЛЬ              СТАТУС              ДАТА ОТПРАВКИ\n")
            self.russia_post_table.insert('end', "-" * 100 + "\n")
            
            if not returns:
                self.russia_post_table.insert('end', "\n(Нет возвратов)\n")
                self.log_russia_post("⚠ Возвратов не найдено")
            else:
                for shipment in returns:
                    tracking = shipment.get('tracking_number', 'N/A').ljust(20)
                    recipient = (shipment.get('recipient_name', 'N/A')[:20]).ljust(24)
                    status = (shipment.get('status', 'N/A')[:16]).ljust(20)
                    sent_date = shipment.get('sent_date', 'N/A').ljust(15)
                    
                    self.russia_post_table.insert('end', f"{tracking}{recipient}{status}{sent_date}\n")
                
                self.log_russia_post(f"✓ Найдено возвратов: {len(returns)}")
            
            self.russia_post_table.config(state=tk.DISABLED)
            
        except Exception as e:
            self.log_russia_post(f"✗ Ошибка при фильтрации: {str(e)}")

    def add_new_shipment_dialog(self):
        dialog = tk.Toplevel(self.root)
        dialog.title("Добавить новое отправление")
        dialog.geometry("500x300")
        dialog.configure(bg=self.colors['bg'])
        dialog.transient(self.root)
        dialog.grab_set()
        
        tk.Label(dialog, text="Трекинг-номер (ШПИ):", bg=self.colors['bg'], font=('Tahoma', 9)).grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        tracking_entry = tk.Entry(dialog, width=30, font=('Tahoma', 9))
        tracking_entry.grid(row=0, column=1, padx=10, pady=10)
        
        tk.Label(dialog, text="Получатель:", bg=self.colors['bg'], font=('Tahoma', 9)).grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        recipient_entry = tk.Entry(dialog, width=30, font=('Tahoma', 9))
        recipient_entry.grid(row=1, column=1, padx=10, pady=10)
        
        tk.Label(dialog, text="Адрес:", bg=self.colors['bg'], font=('Tahoma', 9)).grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)
        address_entry = tk.Entry(dialog, width=30, font=('Tahoma', 9))
        address_entry.grid(row=2, column=1, padx=10, pady=10)
        
        tk.Label(dialog, text="Дата отправки (ДД.ММ.ГГГГ):", bg=self.colors['bg'], font=('Tahoma', 9)).grid(row=3, column=0, padx=10, pady=10, sticky=tk.W)
        date_entry = tk.Entry(dialog, width=30, font=('Tahoma', 9))
        date_entry.grid(row=3, column=1, padx=10, pady=10)
        
        from datetime import datetime
        date_entry.insert(0, datetime.now().strftime('%d.%m.%Y'))
        
        def save_shipment():
            tracking_number = tracking_entry.get().strip()
            recipient_name = recipient_entry.get().strip()
            address = address_entry.get().strip()
            sent_date = date_entry.get().strip()
            
            if not tracking_number:
                messagebox.showerror("Ошибка", "Трекинг-номер обязателен для заполнения", parent=dialog)
                return
            
            existing = self.db.get_shipment_by_tracking(tracking_number)
            if existing:
                messagebox.showerror("Ошибка", f"Отправление с номером {tracking_number} уже существует в базе", parent=dialog)
                return
            
            shipment_id = self.db.add_shipment(
                tracking_number=tracking_number,
                recipient_name=recipient_name,
                address=address,
                status='Отправлено',
                sent_date=sent_date
            )
            
            if shipment_id:
                self.log_russia_post(f"✓ Добавлено отправление: {tracking_number}")
                messagebox.showinfo("Успех", f"Отправление {tracking_number} добавлено в базу", parent=dialog)
                dialog.destroy()
                self.load_shipments_table()
            else:
                self.log_russia_post(f"✗ Ошибка добавления: {tracking_number}")
                messagebox.showerror("Ошибка", "Не удалось добавить отправление в базу данных", parent=dialog)
        
        button_frame = tk.Frame(dialog, bg=self.colors['bg'])
        button_frame.grid(row=4, column=0, columnspan=2, pady=20)
        
        tk.Button(button_frame, text="Сохранить", command=save_shipment, 
                 bg=self.colors['button'], fg=self.colors['text'],
                 font=('Tahoma', 9, 'bold'), padx=20, pady=8).grid(row=0, column=0, padx=5)
        
        tk.Button(button_frame, text="Отмена", command=dialog.destroy,
                 bg=self.colors['button'], fg=self.colors['text'],
                 font=('Tahoma', 9, 'bold'), padx=20, pady=8).grid(row=0, column=1, padx=5)

    def log_scanner(self, message):
        if hasattr(self, 'scanner_log'):
            self.scanner_log.config(state=tk.NORMAL)
            self.scanner_log.insert(tk.END, f"[{self.get_timestamp()}] {message}\n")
            self.scanner_log.see(tk.END)
            self.scanner_log.config(state=tk.DISABLED)
    
    def check_scanner_system(self):
        """Check scanner system requirements and update status indicator."""
        def _check_thread():
            try:
                self.root.after(0, lambda: self.log_scanner("🔍 Проверка системных требований для сканера..."))
                self.root.after(0, lambda: self.scanner_status_label.config(text="🟡 Проверка..."))
                
                # Ensure scanner service is initialized
                if not self.services_initialized['scanner']:
                    self.root.after(0, lambda: self.init_scanner())
                
                # Run system requirements check
                checks = self.scanner.check_system_requirements()
                
                # Update status based on results
                if checks['overall_status'] == 'ready':
                    self.root.after(0, lambda: self.scanner_status_label.config(text="🟢 Готов"))
                    self.root.after(0, lambda: self.log_scanner("✓ Все системные требования выполнены"))
                elif checks['overall_status'] == 'partial':
                    self.root.after(0, lambda: self.scanner_status_label.config(text="🟡 Частично"))
                    self.root.after(0, lambda: self.log_scanner("⚠ Частичная готовность системы"))
                else:
                    self.root.after(0, lambda: self.scanner_status_label.config(text="🔴 Не готов"))
                    self.root.after(0, lambda: self.log_scanner("✗ Система не готова для работы сканера"))
                
                # Log detailed results
                self.root.after(0, lambda: self.log_scanner(""))
                self.root.after(0, lambda: self.log_scanner("Детали проверки:"))
                
                status_icon = lambda x: "✓" if x else "✗"
                
                self.root.after(0, lambda: self.log_scanner(f"  {status_icon(checks['tesseract_installed'])} Tesseract OCR"))
                self.root.after(0, lambda: self.log_scanner(f"  {status_icon(checks['pillow_available'])} Pillow (PIL)"))
                self.root.after(0, lambda: self.log_scanner(f"  {status_icon(checks['pytesseract_available'])} pytesseract"))
                self.root.after(0, lambda: self.log_scanner(f"  {status_icon(checks['twain_available'])} TWAIN драйвер"))
                self.root.after(0, lambda: self.log_scanner(f"  {status_icon(checks['wia_available'])} WIA (Windows)"))
                
                self.root.after(0, lambda: self.log_scanner(""))
                if checks['overall_status'] != 'ready':
                    self.root.after(0, lambda: self.log_scanner("Для полной работы установите:"))
                    if not checks['tesseract_installed']:
                        self.root.after(0, lambda: self.log_scanner("  • Tesseract OCR: https://github.com/tesseract-ocr/tesseract"))
                    if not checks['pillow_available']:
                        self.root.after(0, lambda: self.log_scanner("  • Pillow: pip install Pillow"))
                    if not checks['pytesseract_available']:
                        self.root.after(0, lambda: self.log_scanner("  • pytesseract: pip install pytesseract"))
                    if not checks['twain_available'] and not checks['wia_available']:
                        self.root.after(0, lambda: self.log_scanner("  • Драйверы Canon DR-M260 (TWAIN/WIA)"))
                
            except Exception as e:
                error_msg = str(e)
                self.root.after(0, lambda: self.scanner_status_label.config(text="🔴 Ошибка"))
                self.root.after(0, lambda: self.log_scanner(f"✗ Ошибка проверки: {error_msg}"))
        
        threading.Thread(target=_check_thread, daemon=True).start()

    def launch_scanner_setup(self):
        """Launch scanner setup batch script (Windows only)."""
        self.log_scanner("🚀 Запуск мастера установки компонентов сканера...")
        
        # Check if running on Windows
        if os.name != 'nt':
            self.log_scanner("✗ Установщик компонентов доступен только на Windows")
            messagebox.showwarning(
                "Недоступно",
                "Мастер установки компонентов сканера работает только на Windows.\n\n"
                "Текущая ОС не поддерживается для автоматической установки.\n"
                "Пожалуйста, обратитесь к файлу SCANNER_SETUP_RU.txt для ручной установки.",
                parent=self.root
            )
            return
        
        # Resolve batch path using get_resource_path (works for both bundled and source modes)
        batch_path = get_resource_path('setup_scanner.bat')
        
        if not os.path.exists(batch_path):
            self.log_scanner(f"✗ Файл setup_scanner.bat не найден: {batch_path}")
            messagebox.showerror(
                "Ошибка",
                f"Файл установщика не найден:\n{batch_path}\n\n"
                "Убедитесь что setup_scanner.bat находится в той же папке что и программа.",
                parent=self.root
            )
            return
        
        try:
            # Launch batch file in new console window using os.startfile (Windows only)
            os.startfile(batch_path)
            self.log_scanner("✓ Мастер установки запущен в новом окне консоли")
            self.log_scanner("  После завершения установки компонентов нажмите '🔍 Проверить систему'")
        except Exception as e:
            self.log_scanner(f"✗ Ошибка запуска мастера установки: {str(e)}")
            messagebox.showerror(
                "Ошибка запуска",
                f"Не удалось запустить мастер установки:\n{str(e)}",
                parent=self.root
            )

    def log_mosedo(self, message):
        if hasattr(self, 'mosedo_log'):
            self.mosedo_log.config(state=tk.NORMAL)
            self.mosedo_log.insert(tk.END, f"[{self.get_timestamp()}] {message}\n")
            self.mosedo_log.see(tk.END)
            self.mosedo_log.config(state=tk.DISABLED)
    
    def load_workflows_list(self):
        """Load all workflows from database into listbox."""
        if not self.mosedo_automation:
            return
        
        self.workflows_listbox.delete(0, tk.END)
        workflows = self.db.get_all_workflows(active_only=True)
        
        for wf in workflows:
            steps_count = len(json.loads(wf['steps_json'])) if wf['steps_json'] else 0
            display_text = f"{wf['name']} ({steps_count} шагов)"
            self.workflows_listbox.insert(tk.END, display_text)
            self.workflows_listbox.itemconfig(tk.END, {'id': wf['id']})
        
        self.log_mosedo(f"✓ Загружено workflows: {len(workflows)}")
    
    def on_workflow_select(self, event):
        """Handle workflow selection from listbox."""
        selection = self.workflows_listbox.curselection()
        if not selection:
            return
        
        # Ensure MOSEDO automation is initialized
        if not self.mosedo_automation:
            try:
                self.init_mosedo()
            except Exception as e:
                logger.error(f"Failed to initialize MOSEDO automation: {e}")
                self.log_mosedo(f"✗ Ошибка инициализации MOSEDO: {str(e)}")
                return
        
        if not self.mosedo_automation:
            self.log_mosedo("✗ Не удалось инициализировать MOSEDO automation")
            return
        
        idx = selection[0]
        item_text = self.workflows_listbox.get(idx)
        
        workflows = self.db.get_all_workflows(active_only=True)
        if idx < len(workflows):
            wf = workflows[idx]
            self.current_workflow_id = wf['id']
            self.current_workflow_name = wf['name']
            
            steps = self.mosedo_automation.load_workflow(wf['id'])
            if steps:
                self.current_steps = steps
                self.refresh_steps_tree()
                self.update_workflow_info()
                self.log_mosedo(f"✓ Загружен workflow: {wf['name']}")
            else:
                self.log_mosedo(f"✗ Ошибка загрузки workflow {wf['id']}")
    
    def create_new_workflow(self):
        """Create a new workflow."""
        from tkinter import simpledialog
        
        name = simpledialog.askstring(
            "Новый Workflow",
            "Введите название workflow:",
            parent=self.root
        )
        
        if not name:
            return
        
        self.current_workflow_id = None
        self.current_workflow_name = name
        self.current_steps = []
        self.refresh_steps_tree()
        self.update_workflow_info()
        self.log_mosedo(f"✓ Создан новый workflow: {name}")
    
    def rename_workflow(self):
        """Rename selected workflow."""
        if not self.current_workflow_id:
            messagebox.showwarning("Нет выбора", "Выберите workflow для переименования", parent=self.root)
            return
        
        from tkinter import simpledialog
        new_name = simpledialog.askstring(
            "Переименовать Workflow",
            f"Новое название для '{self.current_workflow_name}':",
            initialvalue=self.current_workflow_name,
            parent=self.root
        )
        
        if new_name and new_name != self.current_workflow_name:
            success = self.db.update_workflow(self.current_workflow_id, name=new_name)
            if success:
                self.current_workflow_name = new_name
                self.load_workflows_list()
                self.update_workflow_info()
                self.log_mosedo(f"✓ Workflow переименован: {new_name}")
            else:
                messagebox.showerror("Ошибка", "Не удалось переименовать workflow", parent=self.root)
    
    def delete_workflow(self):
        """Delete selected workflow."""
        if not self.current_workflow_id:
            messagebox.showwarning("Нет выбора", "Выберите workflow для удаления", parent=self.root)
            return
        
        response = messagebox.askyesno(
            "Подтверждение удаления",
            f"Удалить workflow '{self.current_workflow_name}'?\n\nЭто действие нельзя отменить.",
            parent=self.root
        )
        
        if response:
            success = self.db.delete_workflow(self.current_workflow_id)
            if success:
                self.log_mosedo(f"✓ Workflow удален: {self.current_workflow_name}")
                self.current_workflow_id = None
                self.current_workflow_name = None
                self.current_steps = []
                self.refresh_steps_tree()
                self.update_workflow_info()
                self.load_workflows_list()
            else:
                messagebox.showerror("Ошибка", "Не удалось удалить workflow", parent=self.root)
    
    def add_step(self):
        """Add a new step to current workflow."""
        action = self.action_var.get()
        selector = self.selector_entry.get().strip()
        value = self.value_entry.get().strip()
        description = self.description_entry.get().strip()
        
        if not self.validate_step_input(action, selector, value):
            return
        
        from src.mosedo_automation import WorkflowStep
        step = WorkflowStep(
            action=action,
            selector=selector if selector else None,
            value=value if value else None,
            description=description if description else None
        )
        
        self.current_steps.append(step)
        self.refresh_steps_tree()
        self.update_workflow_info()
        
        self.selector_entry.delete(0, tk.END)
        self.value_entry.delete(0, tk.END)
        self.description_entry.delete(0, tk.END)
        
        self.log_mosedo(f"✓ Добавлен шаг: {action}")
    
    def edit_step(self):
        """Edit selected step."""
        selection = self.steps_tree.selection()
        if not selection:
            messagebox.showwarning("Нет выбора", "Выберите шаг для редактирования", parent=self.root)
            return
        
        idx = self.steps_tree.index(selection[0])
        if idx >= len(self.current_steps):
            return
        
        action = self.action_var.get()
        selector = self.selector_entry.get().strip()
        value = self.value_entry.get().strip()
        description = self.description_entry.get().strip()
        
        if not self.validate_step_input(action, selector, value):
            return
        
        step = self.current_steps[idx]
        step.action = action
        step.selector = selector if selector else None
        step.value = value if value else None
        step.description = description if description else None
        
        self.refresh_steps_tree()
        self.log_mosedo(f"✓ Шаг {idx+1} изменен")
    
    def delete_step(self):
        """Delete selected step."""
        selection = self.steps_tree.selection()
        if not selection:
            messagebox.showwarning("Нет выбора", "Выберите шаг для удаления", parent=self.root)
            return
        
        idx = self.steps_tree.index(selection[0])
        if idx < len(self.current_steps):
            del self.current_steps[idx]
            self.refresh_steps_tree()
            self.update_workflow_info()
            self.log_mosedo(f"✓ Шаг {idx+1} удален")
    
    def move_step_up(self):
        """Move selected step up."""
        selection = self.steps_tree.selection()
        if not selection:
            return
        
        idx = self.steps_tree.index(selection[0])
        if idx > 0 and idx < len(self.current_steps):
            self.current_steps[idx], self.current_steps[idx-1] = self.current_steps[idx-1], self.current_steps[idx]
            self.refresh_steps_tree()
            self.log_mosedo(f"✓ Шаг {idx+1} перемещен вверх")
    
    def move_step_down(self):
        """Move selected step down."""
        selection = self.steps_tree.selection()
        if not selection:
            return
        
        idx = self.steps_tree.index(selection[0])
        if idx < len(self.current_steps) - 1:
            self.current_steps[idx], self.current_steps[idx+1] = self.current_steps[idx+1], self.current_steps[idx]
            self.refresh_steps_tree()
            self.log_mosedo(f"✓ Шаг {idx+1} перемещен вниз")
    
    def save_current_workflow(self):
        """Save current workflow to database."""
        if not self.current_workflow_name:
            messagebox.showwarning("Нет workflow", "Создайте новый workflow или выберите существующий", parent=self.root)
            return
        
        if not self.current_steps:
            messagebox.showwarning("Пустой workflow", "Добавьте хотя бы один шаг перед сохранением", parent=self.root)
            return
        
        import json
        steps_json = json.dumps([step.to_dict() for step in self.current_steps], ensure_ascii=False)
        
        if self.current_workflow_id:
            success = self.db.update_workflow(
                workflow_id=self.current_workflow_id,
                steps_json=steps_json
            )
            if success:
                self.log_mosedo(f"✓ Workflow сохранен: {self.current_workflow_name}")
                self.load_workflows_list()
                messagebox.showinfo("Успешно", f"Workflow '{self.current_workflow_name}' обновлен", parent=self.root)
            else:
                messagebox.showerror("Ошибка", "Не удалось сохранить workflow", parent=self.root)
        else:
            wf_id = self.db.add_workflow(
                name=self.current_workflow_name,
                steps_json=steps_json
            )
            if wf_id:
                self.current_workflow_id = wf_id
                self.log_mosedo(f"✓ Workflow создан (ID: {wf_id}): {self.current_workflow_name}")
                self.load_workflows_list()
                messagebox.showinfo("Успешно", f"Workflow '{self.current_workflow_name}' сохранен", parent=self.root)
            else:
                messagebox.showerror("Ошибка", "Не удалось создать workflow (возможно, название уже существует)", parent=self.root)
    
    def run_workflow(self):
        """Run selected workflow with auto-save."""
        if not self.current_workflow_name:
            messagebox.showwarning("Нет workflow", "Создайте или выберите workflow для запуска", parent=self.root)
            return
        
        if not self.current_steps:
            messagebox.showwarning("Пустой workflow", "Добавьте шаги перед запуском", parent=self.root)
            return
        
        # Ensure MOSEDO automation is initialized
        if not self.mosedo_automation:
            try:
                self.init_mosedo()
            except Exception as e:
                logger.error(f"Failed to initialize MOSEDO automation: {e}")
                messagebox.showerror("MOSEDO недоступен", 
                                   f"Не удалось инициализировать MOSEDO автоматизацию:\n{str(e)}",
                                   parent=self.root)
                return
        
        if not self.mosedo_automation:
            messagebox.showerror("MOSEDO недоступен", 
                               "Не удалось инициализировать MOSEDO автоматизацию.\n"
                               "Проверьте наличие Selenium и ChromeDriver.",
                               parent=self.root)
            return
        
        if self.current_workflow_id:
            import json
            steps_json = json.dumps([step.to_dict() for step in self.current_steps], ensure_ascii=False)
            success = self.db.update_workflow(
                workflow_id=self.current_workflow_id,
                steps_json=steps_json
            )
            if not success:
                messagebox.showerror("Ошибка сохранения", "Не удалось сохранить изменения перед запуском", parent=self.root)
                return
            self.log_mosedo("✓ Изменения сохранены перед запуском")
        else:
            import json
            steps_json = json.dumps([step.to_dict() for step in self.current_steps], ensure_ascii=False)
            wf_id = self.db.add_workflow(
                name=self.current_workflow_name,
                steps_json=steps_json
            )
            if not wf_id:
                messagebox.showerror("Ошибка сохранения", "Сохраните workflow перед запуском", parent=self.root)
                return
            self.current_workflow_id = wf_id
            self.load_workflows_list()
            self.log_mosedo(f"✓ Workflow сохранен (ID: {wf_id}) перед запуском")
        
        self.log_mosedo(f"▶ Запуск workflow: {self.current_workflow_name}")
        self.workflow_status_label.config(text="Статус: Выполняется...")
        
        def _run_thread():
            try:
                success, message = self.mosedo_automation.execute_workflow(
                    self.current_workflow_id,
                    headless=False
                )
                msg_success = message
                self.root.after(0, lambda: self.log_mosedo(f"✓ {msg_success}" if success else f"✗ {msg_success}"))
                self.root.after(0, lambda: self.workflow_status_label.config(
                    text=f"Статус: {'Завершено' if success else 'Ошибка'}"
                ))
            except Exception as e:
                error_msg = str(e)
                self.root.after(0, lambda: self.log_mosedo(f"✗ Ошибка выполнения: {error_msg}"))
                self.root.after(0, lambda: self.workflow_status_label.config(text="Статус: Ошибка"))
        
        threading.Thread(target=_run_thread, daemon=True).start()
    
    def stop_workflow(self):
        """Stop workflow execution."""
        self.log_mosedo("⏹ Остановка workflow...")
        self.workflow_status_label.config(text="Статус: Остановлен")
    
    def start_recording_workflow(self):
        """Start browser recording workflow (automatic action capture)."""
        if not self.mosedo_automation:
            try:
                self.init_mosedo()
            except Exception as e:
                logger.error(f"Failed to initialize MOSEDO automation: {e}")
                messagebox.showerror("MOSEDO недоступен", 
                                   f"Не удалось инициализировать MOSEDO автоматизацию:\n{str(e)}\n\n"
                                   "Проверьте наличие Selenium и ChromeDriver.",
                                   parent=self.root)
                return
        
        if not self.mosedo_automation:
            messagebox.showwarning("MOSEDO недоступен", 
                                 "Не удалось инициализировать MOSEDO автоматизацию.\n"
                                 "Проверьте наличие Selenium и ChromeDriver.",
                                 parent=self.root)
            return
        
        from tkinter import simpledialog
        
        # Ask for workflow name
        name = simpledialog.askstring(
            "Запись workflow",
            "Введите название workflow для записи:",
            parent=self.root
        )
        
        if not name:
            return
        
        # Ask for start URL
        start_url = simpledialog.askstring(
            "Запись workflow",
            "Введите начальный URL для записи\n(например: https://mosedo.mos.ru):",
            initialvalue="https://mosedo.mos.ru",
            parent=self.root
        )
        
        if not start_url:
            return
        
        # Check if Selenium is available
        from src.mosedo_automation import SELENIUM_AVAILABLE
        if not SELENIUM_AVAILABLE:
            messagebox.showerror(
                "Selenium недоступен",
                "Для автоматической записи необходим Selenium.\n\n"
                "Установите: pip install selenium\n"
                "И убедитесь что Chrome/ChromeDriver доступен.",
                parent=self.root
            )
            self.log_mosedo("✗ Selenium не установлен - запись невозможна")
            return
        
        self.log_mosedo(f"🔴 Начало записи workflow: {name}")
        self.log_mosedo(f"   URL: {start_url}")
        self.log_mosedo("   Откроется браузер - выполняйте действия...")
        self.log_mosedo("   Нажмите '⏹ Остановить запись' когда закончите")
        
        # Start recording in thread to avoid blocking UI
        def _start_recording():
            try:
                success = self.mosedo_automation.recorder.start_recording(name, start_url)
                if success:
                    self.root.after(0, lambda: self.log_mosedo("✓ Браузер открыт, запись началась"))
                    self.root.after(0, lambda: self.log_mosedo("🔴 Все клики и ввод текста записываются автоматически"))
                    self.root.after(0, lambda: self.log_mosedo("⌨️ Нажмите Ctrl+Shift+S в браузере для остановки"))
                    
                    # Set workflow name
                    self.current_workflow_name = name
                    self.current_workflow_id = None  # New workflow, not saved yet
                    self.root.after(0, self.update_workflow_info)
                    
                    # Start real-time UI updates
                    self.root.after(0, self.update_recording_status)
                else:
                    self.root.after(0, lambda: self.log_mosedo("✗ Не удалось начать запись"))
                    self.root.after(0, lambda: messagebox.showerror(
                        "Ошибка записи",
                        "Не удалось запустить браузер для записи.\n"
                        "Проверьте что Chrome и ChromeDriver установлены.",
                        parent=self.root
                    ))
            except Exception as e:
                error_msg = str(e)
                self.root.after(0, lambda: self.log_mosedo(f"✗ Ошибка записи: {error_msg}"))
        
        threading.Thread(target=_start_recording, daemon=True).start()
    
    def update_recording_status(self):
        """Update UI in real-time while recording is in progress."""
        if not self.mosedo_automation or not self.mosedo_automation.recorder:
            return
        
        if not self.mosedo_automation.recorder.is_recording:
            # Recording stopped
            self.workflow_status_label.config(text="Статус: Запись завершена")
            return
        
        try:
            # Get current steps from recorder
            current_steps = self.mosedo_automation.recorder.get_current_steps()
            
            # Update step count
            self.workflow_steps_label.config(text=f"Шагов: {len(current_steps)} 🔴")
            self.workflow_status_label.config(text="Статус: 🔴 Идет запись...")
            
            # Update steps in current_steps for refresh
            self.current_steps = current_steps
            self.refresh_steps_tree()
            
            # Schedule next update
            self.root.after(500, self.update_recording_status)
            
        except Exception as e:
            logger.error(f"Error updating recording status: {e}")
    
    def stop_recording_workflow(self):
        """Stop browser recording and load captured steps."""
        if not self.mosedo_automation or not self.mosedo_automation.recorder:
            messagebox.showwarning("Нет записи", "Запись не запущена", parent=self.root)
            return
        
        if not self.mosedo_automation.recorder.is_recording:
            messagebox.showwarning("Нет записи", "Запись workflow не активна", parent=self.root)
            return
        
        self.log_mosedo("⏹ Остановка записи...")
        
        def _stop_recording():
            try:
                success, steps = self.mosedo_automation.recorder.stop_recording()
                
                if success and steps:
                    self.root.after(0, lambda: self.log_mosedo(f"✓ Запись остановлена, получено шагов: {len(steps)}"))
                    
                    # Convert WorkflowStep objects to our format
                    from src.mosedo_automation import WorkflowStep
                    self.current_steps = steps
                    
                    # Update UI
                    self.root.after(0, self.refresh_steps_tree)
                    self.root.after(0, self.update_workflow_info)
                    
                    self.root.after(0, lambda: messagebox.showinfo(
                        "Запись завершена",
                        f"Записано {len(steps)} действий.\n\n"
                        f"Проверьте список шагов и нажмите '💾 Сохранить' для сохранения workflow.",
                        parent=self.root
                    ))
                else:
                    self.root.after(0, lambda: self.log_mosedo("✗ Не удалось остановить запись или нет шагов"))
                    
            except Exception as e:
                error_msg = str(e)
                self.root.after(0, lambda: self.log_mosedo(f"✗ Ошибка остановки записи: {error_msg}"))
        
        threading.Thread(target=_stop_recording, daemon=True).start()
    
    def validate_step_input(self, action, selector, value):
        """Validate step input fields."""
        if action in ['click', 'type'] and not selector:
            messagebox.showwarning("Ошибка валидации", f"Action '{action}' требует Selector (CSS)", parent=self.root)
            return False
        
        if action == 'type' and not value:
            messagebox.showwarning("Ошибка валидации", "Action 'type' требует Value (текст для ввода)", parent=self.root)
            return False
        
        if action == 'navigate' and not value:
            messagebox.showwarning("Ошибка валидации", "Action 'navigate' требует Value (URL)", parent=self.root)
            return False
        
        return True
    
    def refresh_steps_tree(self):
        """Refresh steps treeview with current steps."""
        for item in self.steps_tree.get_children():
            self.steps_tree.delete(item)
        
        for i, step in enumerate(self.current_steps):
            self.steps_tree.insert('', tk.END, values=(
                i + 1,
                step.action,
                step.selector or '',
                step.value or '',
                step.description or ''
            ))
    
    def update_workflow_info(self):
        """Update workflow info labels."""
        if self.current_workflow_name:
            self.workflow_name_label.config(text=f"Название: {self.current_workflow_name}")
            self.workflow_steps_label.config(text=f"Шагов: {len(self.current_steps)}")
            status = "Сохранен" if self.current_workflow_id else "Не сохранен"
            self.workflow_status_label.config(text=f"Статус: {status}")
        else:
            self.workflow_name_label.config(text="Название: ---")
            self.workflow_steps_label.config(text="Шагов: 0")
            self.workflow_status_label.config(text="Статус: Не выбран")

    def get_timestamp(self):
        from datetime import datetime
        return datetime.now().strftime('%H:%M:%S')

    def on_drop(self, event):
        files = self.root.tk.splitlist(event.data)
        for file in files:
            if file.lower().endswith('.pdf'):
                self.add_file(file)

    def update_status(self):
        try:
            template_count = len(list(Path("templates").glob("*.docx")))
            conn = self.db._get_connection()
            cursor = conn.cursor()
            cursor.execute('SELECT COUNT(*) FROM directors WHERE is_active = 1')
            director_count = cursor.fetchone()[0]
            cursor.execute('SELECT COUNT(*) FROM inspectors_chiefs WHERE is_active = 1')
            inspector_chiefs_count = cursor.fetchone()[0]
            conn.close()
            self.status_label.config(
                text=f"[SUCCESS] Готово к работе | Шаблонов: {template_count} | Директоров: {director_count} | Начальников инспекций: {inspector_chiefs_count}"
            )
        except sqlite3.Error as e:
            self.log(f"[ERROR] Ошибка базы данных: {str(e)}")
            self.status_label.config(text="[ERROR] Ошибка базы данных")
        except Exception as e:
            self.log(f"[ERROR] Ошибка обновления статуса: {str(e)}")
            self.status_label.config(text="[WARNING] Проверка настроек...")

    def select_files(self):
        files = filedialog.askopenfilenames(
            title="Выберите PDF файлы",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        for file in files:
            self.add_file(file)

    def add_file(self, filepath):
        if filepath not in self.processing_queue:
            self.processing_queue.append(filepath)
            self.document_outputs.append(None)
            self.refresh_listbox()
            self.log(f"Добавлен файл: {os.path.basename(filepath)}")

    def clear_files(self):
        self.processing_queue.clear()
        self.document_outputs.clear()
        self.refresh_listbox()
        self.update_preview("")
        self.log("Очередь обработки очищена (история сохранена)")

    def extract_docx_text(self, docx_path):
        try:
            from docx import Document
            doc = Document(docx_path)
            
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    if paragraph.style.name.startswith('Heading'):
                        text_parts.append(f"\n{'='*60}\n{text}\n{'='*60}\n")
                    else:
                        text_parts.append(text)
            
            for table in doc.tables:
                text_parts.append("\n[ТАБЛИЦА]\n")
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
                text_parts.append("[КОНЕЦ ТАБЛИЦЫ]\n")
            
            return "\n".join(text_parts)
        except Exception as e:
            return f"Ошибка при извлечении текста: {str(e)}"

    def update_preview(self, text):
        self.preview_text.config(state=tk.NORMAL)
        self.preview_text.delete(1.0, tk.END)
        if text:
            self.preview_text.insert(1.0, text)
        else:
            self.preview_text.insert(1.0, "Предпросмотр документа будет отображаться здесь после обработки PDF файлов.")
        self.preview_text.config(state=tk.DISABLED)

    def refresh_listbox(self):
        self.files_listbox.delete(0, tk.END)
        self.listbox_index_map = []
        
        if self.processing_queue:
            for i, pdf_path in enumerate(self.processing_queue):
                self.files_listbox.insert(tk.END, f"▶ {os.path.basename(pdf_path)}")
                self.listbox_index_map.append({
                    'type': 'queue',
                    'queue_index': i
                })
        
        for run_index, run in enumerate(reversed(self.processing_runs)):
            actual_run_index = len(self.processing_runs) - 1 - run_index
            
            self.files_listbox.insert(tk.END, f"═══ {run['timestamp']} ═══")
            self.listbox_index_map.append({
                'type': 'header',
                'run_index': actual_run_index
            })
            
            for entry_index, entry in enumerate(run['results']):
                pdf_name = os.path.basename(entry['pdf_path'])
                self.files_listbox.insert(tk.END, f"  ✓ {pdf_name}")
                self.listbox_index_map.append({
                    'type': 'history',
                    'run_index': actual_run_index,
                    'entry_index': entry_index
                })

    def on_pdf_select(self, event):
        selection = self.files_listbox.curselection()
        if not selection:
            return
        
        idx = selection[0]
        if idx >= len(self.listbox_index_map):
            return
        
        map_entry = self.listbox_index_map[idx]
        
        if map_entry['type'] == 'header':
            return
        
        if map_entry['type'] == 'queue':
            queue_index = map_entry['queue_index']
            if queue_index < len(self.document_outputs):
                doc_info = self.document_outputs[queue_index]
                if doc_info:
                    self.update_preview(doc_info['preview_text'])
                else:
                    self.update_preview("Документ ещё не сгенерирован для этого PDF файла.")
        elif map_entry['type'] == 'history':
            run_index = map_entry['run_index']
            entry_index = map_entry['entry_index']
            if run_index < len(self.processing_runs):
                run = self.processing_runs[run_index]
                if entry_index < len(run['results']):
                    entry = run['results'][entry_index]
                    self.update_preview(entry['preview_text'])

    def log(self, message):
        self.log_text.config(state=tk.NORMAL)
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.log_text.config(state=tk.DISABLED)

    def ezp_log(self, message):
        self.ezp_log_text.config(state=tk.NORMAL)
        self.ezp_log_text.insert(tk.END, message + "\n")
        self.ezp_log_text.see(tk.END)
        self.ezp_log_text.config(state=tk.DISABLED)


    def load_pdf_files(self):
        """Load PDF files using file dialog."""
        file_paths = filedialog.askopenfilenames(
            title="Выберите PDF файлы",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        if file_paths:
            self.ezp_pdf_files = list(file_paths)
            self.ezp_log(f"✓ Выбрано {len(self.ezp_pdf_files)} PDF файлов")
            
            # Update drop zone
            self.ezp_drop_zone.config(
                text=f"✓ Загружено {len(self.ezp_pdf_files)} PDF файлов\n\n"
                     f"Готово к обработке",
                bg='#d5f4e6',
                fg='#27ae60'
            )
    
    def ezp_on_drop(self, event):
        """Handle drag & drop of PDF files."""
        files = self.root.tk.splitlist(event.data)
        
        # Collect all PDF files (expand folders if dropped)
        pdf_files = []
        for item in files:
            if os.path.isdir(item):
                # Folder dropped - collect all PDFs from it
                for filename in os.listdir(item):
                    if filename.lower().endswith('.pdf'):
                        pdf_files.append(os.path.join(item, filename))
            elif item.lower().endswith('.pdf'):
                # PDF file dropped directly
                pdf_files.append(item)
        
        if not pdf_files:
            messagebox.showwarning("Нет PDF файлов", 
                                  "Перетащенные файлы не содержат PDF документов",
                                  parent=self.root)
            return
        
        # Store the list of PDF files
        self.ezp_pdf_files = pdf_files
        self.ezp_log(f"✓ Перетащено {len(pdf_files)} PDF файлов")
        
        self.ezp_drop_zone.config(
            text=f"✓ Загружено {len(pdf_files)} PDF файлов\n\n"
                 f"Готово к обработке",
            bg='#d5f4e6',
            fg='#27ae60'
        )
    
    def process_ezp_batch(self):
        """Process EZP with automatic template loading."""
        if not self.ezp_pdf_files:
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите или перетащите PDF файлы", parent=self.root)
            return
        
        # Start batch processing in thread
        def _batch_process():
            try:
                from src.ezp_processor import EZPProcessor
                
                self.ezp_log("\n" + "="*70)
                self.ezp_log("🚀 НАЧАЛО ОБРАБОТКИ ЭЗП")
                self.ezp_log("="*70)
                self.ezp_log("📋 Использую шаблон: templates/ezp_template.xls")
                self.ezp_log(f"📄 Файлов к обработке: {len(self.ezp_pdf_files)}")
                self.ezp_log("")
                
                # Initialize processor
                processor = EZPProcessor()
                
                # Callback for progress
                def update_progress(value):
                    if self.ezp_progress:
                        self.root.after(0, lambda v=value: self.ezp_progress.config(value=v))
                    if self.ezp_progress_label:
                        self.root.after(0, lambda v=value: self.ezp_progress_label.config(
                            text=f"Прогресс: {v}%"
                        ))
                
                # Callback for logs
                def log_message(msg):
                    self.root.after(0, lambda m=msg: self.ezp_log(m))
                
                # Process EZP with automatic template loading
                success, message, stats = processor.process_ezp(
                    pdf_files=self.ezp_pdf_files,
                    progress_callback=update_progress,
                    log_callback=log_message
                )
                
                # Complete
                if success:
                    if self.ezp_progress:
                        self.root.after(0, lambda: self.ezp_progress.config(value=100))
                    if self.ezp_progress_label:
                        self.root.after(0, lambda: self.ezp_progress_label.config(
                            text="Завершено!"
                        ))
                    
                    # Show completion message
                    self.root.after(0, lambda m=message: messagebox.showinfo(
                        "Обработка завершена",
                        f"{m}\n\nФайл сохранен в папке output/",
                        parent=self.root
                    ))
                else:
                    self.root.after(0, lambda m=message: messagebox.showerror(
                        "Ошибка",
                        f"Ошибка обработки ЭЗП:\n{m}",
                        parent=self.root
                    ))
                
            except Exception as e:
                logger.error(f"Error in batch processing: {e}")
                self.root.after(0, lambda err=str(e): self.ezp_log(f"\n✗ КРИТИЧЕСКАЯ ОШИБКА: {err}"))
                self.root.after(0, lambda err=str(e): messagebox.showerror(
                    "Ошибка",
                    f"Ошибка batch обработки:\n{err}",
                    parent=self.root
                ))
        
        threading.Thread(target=_batch_process, daemon=True).start()
    
    def _detect_recipient_type_ai(self, address: str) -> str:
        """Use AI to detect if recipient is legal entity or individual."""
        if not self.ai_assistant or not self.ai_assistant.is_online:
            return "физическое лицо"
        
        try:
            # Extract organization name from address (usually first line)
            lines = address.split('\n')
            org_name = lines[0] if lines else address
            
            # Ask AI to classify
            prompt = f"""Определи тип получателя по названию организации или имени:

"{org_name}"

Ответь ОДНИМ словом:
- "юридическое" если это организация (ООО, АО, ГБУ, ФГБУ, Департамент, Префектура и т.д.)
- "физическое" если это частное лицо (Ф.И.О.)

Ответ:"""
            
            response = self.ai_assistant.generate_text(prompt, max_tokens=10)
            
            if "юридическое" in response.lower() or "юр" in response.lower():
                return "юридическое лицо"
            else:
                return "физическое лицо"
                
        except Exception as e:
            logger.error(f"AI detection error: {e}")
            return "физическое лицо"  # Fallback

    def process_ezp(self):
        """Legacy function - redirects to process_ezp_batch"""
        self.process_ezp_batch()
    
    def process_ezp_thread(self):
        """Legacy function - no longer used, kept for compatibility"""
        pass

    def process_files(self):
        if not self.processing_queue:
            messagebox.showwarning("Предупреждение", "Пожалуйста, выберите PDF файлы для обработки")
            return
        
        self.document_outputs = [None] * len(self.processing_queue)

        thread = threading.Thread(target=self.process_files_thread)
        thread.daemon = True
        thread.start()

    def process_files_thread(self):
        import datetime
        
        total_files = len(self.processing_queue)
        success_count = 0
        error_count = 0
        
        timestamp = datetime.datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        current_run = {
            'timestamp': timestamp,
            'results': []
        }

        for index, pdf_file in enumerate(self.processing_queue):
            try:
                self.log(f"\n{'='*60}")
                self.log(f"Обработка файла {index+1}/{total_files}: {os.path.basename(pdf_file)}")

                # Validate file exists
                if not os.path.exists(pdf_file):
                    raise FileNotFoundError(f"Файл не найден: {pdf_file}")

                self.log("Извлечение данных из PDF...")
                citizen_data = self.parser.parse_pdf(pdf_file, database=self.db)

                # Validate extracted data
                if not citizen_data.get('last_name') or not citizen_data.get('first_name'):
                    raise ValueError("Не удалось извлечь ФИО из документа")

                self.log(f"  ФИО: {citizen_data['full_name']}")
                self.log(f"  Email: {citizen_data['email']}")
                self.log(f"  Номер ОАТИ: {citizen_data['oati_number']}")

                self.log("Склонение имени...")
                declined = self.decliner.decline_full_name(
                    citizen_data['last_name'],
                    citizen_data['first_name'],
                    citizen_data['middle_name']
                )

                short_name_dative = self.decliner.get_short_name_dative(
                    citizen_data['last_name'],
                    citizen_data['first_name'],
                    citizen_data['middle_name']
                )

                salutation = self.decliner.get_full_salutation(
                    citizen_data['first_name'],
                    citizen_data['middle_name'],
                    declined['gender']
                )

                self.log(f"  Короткое ФИО (дательный): {short_name_dative}")
                self.log(f"  Обращение: {salutation}")

                self.log("Формирование параметров для шаблона...")
                portal_source = self.parser.extract_portal_source()

                recipients = self.parser.extract_recipients_from_resolution()
                self.log(f"  Найдено получателей в резолюции: {len(recipients)}")
                for r in recipients:
                    self.log(f"    - {r['last_name']} {r['first_initial']}.{r['middle_initial']}.")

                # Get all inspector chiefs surnames from database for deadline checking
                all_inspectors = self.db.get_all_persons(role='inspector')
                inspector_surnames = []
                for inspector in all_inspectors:
                    # Add both nominative and dative forms for robust matching
                    inspector_surnames.append(inspector['last_name'])
                    if inspector.get('last_name_dative'):
                        inspector_surnames.append(inspector['last_name_dative'])

                # Extract departments from matched directors (not inspectors!)
                matched_departments = []
                has_union_response = False
                inspector_found = False

                for recipient in recipients:
                    match = self.db.match_recipient_from_resolution(
                        recipient['last_name'],
                        recipient['first_initial'],
                        recipient['middle_initial'],
                        self.decliner
                    )
                    if match:
                        self.log(f"  ✓ Найден в БД: {recipient['last_name']} -> {match['type']}")
                        if match['type'] == 'director':
                            # Directors determine recipient departments
                            dept_name = match['data'].get('department', '')
                            if dept_name and dept_name not in matched_departments:
                                if 'объединение административно-технических инспекций' not in dept_name.lower():
                                    matched_departments.append(dept_name)
                                    self.log(f"    → Добавлен департамент: {dept_name}")
                                else:
                                    self.log(f"    → Пропущен (ОАТИ): {dept_name}")
                        elif match['type'] == 'inspector':
                            # Inspector chief found
                            inspector_found = True
                            self.log(f"    → Это инспектор (будет проверка на ответ Объединения)")
                    else:
                        self.log(f"  ✗ НЕ найден в БД: {recipient['last_name']} {recipient['first_initial']}.{recipient['middle_initial']}.")

                # Check if Union response is required (inspector + ОАТИ phrases + NO "НПА")
                has_union_response = self.parser.check_union_response_required(
                    database=self.db
                )

                # Use matched departments if found, otherwise fall back to regex extraction
                # NOTE: "Объединение" is NOT a recipient department - it's filtered out by parser
                if matched_departments:
                    non_union_depts = matched_departments
                    self.log(f"  Департаменты определены через БД: {len(non_union_depts)}")
                else:
                    non_union_depts = citizen_data['departments']  # Already filtered by parser
                    self.log(f"  Департаменты определены через текст: {len(non_union_depts)}")

                # If no departments found at all, warn but continue with default
                if not non_union_depts:
                    self.log("[WARNING] Департаменты не найдены, будет использован дефолтный департамент")
                    non_union_depts = []  # Will use default "компетентный орган" in template

                num_departments = len(non_union_depts)

                # Determine if Union should be added to recipients
                # Union responds if: inspector chief found + has deadline + NO "НПА"
                add_union = has_union_response
                
                # CRITICAL VALIDATION: If Union response needed but NO departments found, warn user
                if add_union and not non_union_depts:
                    self.log("[ERROR] Найдено Объединение (инспектор), но НЕ найдены департаменты-получатели!")
                    self.log("[ERROR] Это НЕПРАВИЛЬНАЯ ситуация - если есть инспектор, должны быть департаменты!")
                    self.log("[ERROR] Проверьте резолюцию вручную и добавьте недостающих руководителей в БД")

                # Build declined departments list (never includes "Объединение")
                declined_depts = [self.decliner.decline_text_to_accusative(dept).replace('в ', '', 1) 
                                 for dept in non_union_depts] if non_union_depts else []

                # AUTOMATIC LAW PART CALCULATION using typed recipient classification
                # Import classify_recipients for automatic law part determination
                from src.recipient_types import classify_recipients
                
                # Classify recipients (departments without Union)
                recipients_for_classification = classify_recipients(non_union_depts)
                
                # Calculate law part automatically based on recipients + Union flag
                # Rule: ч. 3 = exactly 1 recipient WITHOUT Union
                #       ч. 4 = 2+ recipients OR Union involved
                if add_union:
                    # Union (ОАТИ) involved - always use ч. 4
                    actual_law_part = '4'
                    self.log(f"  [AUTO] Часть закона: ч.{actual_law_part} (ОАТИ присутствует)")
                else:
                    # No Union - use automatic calculation
                    actual_law_part = recipients_for_classification.calculate_law_part()
                    self.log(f"  [AUTO] Часть закона: ч.{actual_law_part} (автоматический расчёт)")
                
                # Construct departments list with proper formatting: "в X и в Y"
                if add_union:
                    # Union response needed - prepend full Union name to departments list
                    union_full_name = 'Объединение административно-технических инспекций города Москвы'
                    union_declined = self.decliner.decline_text_to_accusative(union_full_name).replace('в ', '', 1)
                    # Add (далее - Объединение) suffix to FIRST mention only
                    union_with_suffix = union_declined + ' (далее - Объединение)'
                    
                    if declined_depts:
                        # Prepend Union to departments list and format as "в X и в Y"
                        all_recipients = [union_with_suffix] + declined_depts
                        if len(all_recipients) == 1:
                            departments_list = 'в ' + all_recipients[0]
                        elif len(all_recipients) == 2:
                            departments_list = 'в ' + all_recipients[0] + ' и в ' + all_recipients[1]
                        else:
                            # For 3+ organizations: "в X, в Y и в Z"
                            departments_list = ', '.join(['в ' + dept for dept in all_recipients[:-1]]) + ' и в ' + all_recipients[-1]
                    else:
                        # Only Union if no departments found
                        departments_list = 'в ' + union_with_suffix
                    union_suffix = ''  # Already included in departments_list
                else:
                    # No Union - just regular departments formatted as "в X и в Y"
                    if declined_depts:
                        if len(declined_depts) == 1:
                            departments_list = 'в ' + declined_depts[0]
                        elif len(declined_depts) == 2:
                            departments_list = 'в ' + declined_depts[0] + ' и в ' + declined_depts[1]
                        else:
                            # For 3+ departments: "в X, в Y и в Z"
                            departments_list = ', '.join(['в ' + dept for dept in declined_depts[:-1]]) + ' и в ' + declined_depts[-1]
                    else:
                        departments_list = 'в компетентный орган'
                    union_suffix = ''

                self.log(f"  Источник портала: {portal_source}")
                self.log(f"  Департаментов (без Объединения): {num_departments}")
                self.log(f"  Нужен ответ Объединения: {'Да' if add_union else 'Нет'}")
                self.log(f"  Часть закона: ч.{actual_law_part}")
                self.log(f"  Список департаментов: {departments_list}")

                self.log("Генерация Word документа...")
                declined_with_short = declined.copy()
                declined_with_short['full_name'] = short_name_dative

                citizen_data['law_part'] = actual_law_part
                
                # Construct union_paragraph based on whether union response is needed
                union_paragraph = ''
                if add_union:
                    question_text = self.parser.extract_question_text()
                    if question_text:
                        # Try AI-assisted union_paragraph generation first
                        ai_paragraph = None
                        if self.ai_assistant:
                            try:
                                status = self.ai_assistant.get_status()
                                if status['is_online']:
                                    self.log("[AI] Генерация параграфа про Объединение с помощью AI...")
                                    ai_paragraph = self.ai_assistant.generate_union_paragraph(question_text)
                                    if ai_paragraph:
                                        self.log("[AI] ✓ Параграф успешно сгенерирован AI")
                                        union_paragraph = ai_paragraph
                                    else:
                                        self.log("[AI] ⚠ AI вернул пустой результат, используется стандартный текст")
                            except Exception as e:
                                self.log(f"[AI] ✗ Ошибка при генерации AI: {e}, используется стандартный текст")
                        
                        # Fallback: use rule-based generation if AI not available or failed
                        if not ai_paragraph:
                            # Decline question text to accusative case, remove leading 'в'
                            question_declined = self.decliner.decline_text_to_accusative(question_text).replace('в ', '', 1)
                            union_paragraph = f"\n\nДополнительно сообщаем, что вопрос {question_declined} будет рассмотрен Объединением в рамках компетенции в установленные законом сроки."
                            self.log(f"  Вопрос для Объединения: {question_declined}")
                    else:
                        # Fallback if question not found
                        union_paragraph = "\n\nДополнительно сообщаем, что обращение будет рассмотрено Объединением в рамках компетенции в установленные законом сроки."
                        self.log("[WARNING] Вопрос для Объединения не найден в резолюции, используется стандартный текст")
                

                output_file = self.generator.process_citizen_document(
                    citizen_data,
                    declined_with_short,
                    salutation,
                    non_union_depts,
                    portal_source=portal_source,
                    union_suffix=union_suffix,
                    union_paragraph=union_paragraph,
                    departments_list=departments_list,
                    has_inspector=add_union
                )

                # Detailed logging with full paths
                abs_output_path = os.path.abspath(output_file)
                self.log(f"[OK] Документ создан: {os.path.basename(output_file)}")
                self.log(f"[PATH] Полный путь: {abs_output_path}")
                self.log(f"[CHECK] Файл существует: {os.path.exists(abs_output_path)}")
                if os.path.exists(abs_output_path):
                    file_size = os.path.getsize(abs_output_path)
                    self.log(f"[SIZE] Размер файла: {file_size} байт ({file_size/1024:.1f} KB)")
                
                self.log("Извлечение текста для предпросмотра...")
                preview_text = self.extract_docx_text(abs_output_path)
                
                preview_data = {
                    'docx_path': abs_output_path,
                    'preview_text': preview_text
                }
                
                if index >= len(self.document_outputs):
                    self.document_outputs.append(preview_data)
                else:
                    self.document_outputs[index] = preview_data
                
                current_run['results'].append({
                    'pdf_path': pdf_file,
                    'docx_path': abs_output_path,
                    'preview_text': preview_text
                })
                
                self.root.after(0, self.update_preview, preview_text)
                self.log("[OK] Предпросмотр обновлён")

                self.db.add_processing_record(
                    pdf_file,
                    citizen_data['full_name'],
                    citizen_data['oati_number'],
                    citizen_data['portal_id'],
                    output_file,
                    'success'
                )

            except FileNotFoundError as e:
                error_count += 1
                self.log(f"[ERROR] Файл не найден: {str(e)}")
                self.db.add_processing_record(
                    pdf_file, "", "", "", "", f'error: file_not_found - {str(e)}'
                )
            except ValueError as e:
                error_count += 1
                self.log(f"[ERROR] Некорректные данные: {str(e)}")
                self.db.add_processing_record(
                    pdf_file, "", "", "", "", f'error: invalid_data - {str(e)}'
                )
            except Exception as e:
                error_count += 1
                self.log(f"[ERROR] Непредвиденная ошибка: {str(e)}")
                import traceback
                self.log(f"[DEBUG] Traceback: {traceback.format_exc()}")
                self.db.add_processing_record(
                    pdf_file, "", "", "", "", f'error: {str(e)}'
                )
            else:
                success_count += 1

            progress = ((index + 1) / total_files) * 100
            self.progress_var.set(progress)

        if current_run['results']:
            self.processing_runs.append(current_run)
            self.log(f"[HISTORY] Сохранено {len(current_run['results'])} результатов в историю: {timestamp}")
        
        self.root.after(0, self.refresh_listbox)
        
        self.log(f"\n{'='*60}")
        self.log(f"Обработка завершена!")
        self.log(f"Всего файлов: {total_files}")
        self.log(f"Успешно: {success_count}")
        self.log(f"Ошибок: {error_count}")
        self.log(f"Результаты сохранены в папке: output/")

        if error_count > 0:
            messagebox.showwarning("Обработка завершена с ошибками", 
                                  f"Всего файлов: {total_files}\n"
                                  f"Успешно: {success_count}\n"
                                  f"Ошибок: {error_count}\n\n"
                                  f"Подробности в журнале обработки")
        else:
            messagebox.showinfo("Готово", f"Обработка завершена!\nОбработано файлов: {total_files}")

        self.progress_var.set(0)

    def open_departments_window(self):
        dept_window = tk.Toplevel(self.root)
        dept_window.title("Управление директорами")
        dept_window.geometry("1200x600")

        DirectorsManager(dept_window, self.db, self.decliner)

    def open_inspector_chiefs_window(self):
        inspector_window = tk.Toplevel(self.root)
        inspector_window.title("Управление начальниками инспекций")
        inspector_window.geometry("1200x600")

        InspectorChiefsManager(inspector_window, self.db, self.decliner)
    
    def open_organizations_window(self):
        org_window = tk.Toplevel(self.root)
        org_window.title("Управление организациями")
        org_window.geometry("1200x700")

        OrganizationsManager(org_window, self.db, self.root)

    def open_history_window(self):
        history_window = tk.Toplevel(self.root)
        history_window.title("История обработки")
        history_window.geometry("800x600")
        history_window.configure(bg=self.colors['bg'])
        
        main_frame = tk.Frame(history_window, bg=self.colors['bg'], padx=20, pady=20)
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        title_label = tk.Label(main_frame, 
                              text="История обработки документов", 
                              font=('Segoe UI', 14, 'bold'),
                              bg=self.colors['bg'],
                              fg=self.colors['text'])
        title_label.pack(pady=(0, 15))
        
        info_text = tk.Text(main_frame, 
                           bg='white',
                           fg=self.colors['text'],
                           font=('Segoe UI', 10),
                           relief=tk.FLAT,
                           wrap=tk.WORD,
                           padx=15,
                           pady=15)
        info_text.pack(fill=tk.BOTH, expand=True)
        
        if not self.processing_runs:
            info_text.insert('1.0', "История обработки пуста.\n\nОбработанные документы будут отображаться в списке файлов с временными метками.")
        else:
            info_text.insert('1.0', f"Всего запусков обработки: {len(self.processing_runs)}\n\n")
            
            for run_index, run in enumerate(reversed(self.processing_runs)):
                actual_run_index = len(self.processing_runs) - 1 - run_index
                timestamp = run['timestamp']
                results_count = len(run['results'])
                
                info_text.insert(tk.END, f"{'='*60}\n")
                info_text.insert(tk.END, f"Запуск #{actual_run_index + 1} - {timestamp}\n")
                info_text.insert(tk.END, f"Обработано файлов: {results_count}\n")
                info_text.insert(tk.END, f"{'='*60}\n\n")
                
                for entry_index, entry in enumerate(run['results']):
                    pdf_name = os.path.basename(entry['pdf_path'])
                    docx_name = os.path.basename(entry['docx_path'])
                    info_text.insert(tk.END, f"  {entry_index + 1}. {pdf_name}\n")
                    info_text.insert(tk.END, f"     → {docx_name}\n\n")
                
                info_text.insert(tk.END, "\n")
            
            info_text.insert(tk.END, "\nДля просмотра документов выберите их в основном списке файлов.")
        
        info_text.config(state=tk.DISABLED)
        
        close_btn = tk.Button(main_frame, 
                             text="Закрыть",
                             command=history_window.destroy,
                             bg=self.colors['button'],
                             fg=self.colors['text'],
                             font=('Tahoma', 9, 'bold'),
                             relief=tk.RAISED,
                             cursor='hand2',
                             padx=20,
                             pady=8,
                             bd=2)
        close_btn.pack(pady=(15, 0))


class DirectorsManager:
    def __init__(self, window, db, decliner):
        self.window = window
        self.db = db
        self.decliner = decliner

        main_frame = ttk.Frame(window, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        window.columnconfigure(0, weight=1)
        window.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)

        title = ttk.Label(main_frame, text="Справочник директоров департаментов", 
                         font=('Arial', 14, 'bold'))
        title.grid(row=0, column=0, pady=10)

        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, pady=5, sticky=(tk.W, tk.E))

        ttk.Button(button_frame, text="Добавить", command=self.add_director).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Редактировать", command=self.edit_director).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="Редактировать склонения", command=self.edit_declensions).grid(row=0, column=2, padx=5)
        ttk.Button(button_frame, text="Удалить", command=self.delete_director).grid(row=0, column=3, padx=5)
        ttk.Button(button_frame, text="Обновить", command=self.refresh_list).grid(row=0, column=4, padx=5)

        tree_frame = ttk.Frame(main_frame)
        tree_frame.grid(row=2, column=0, pady=10, sticky=(tk.W, tk.E, tk.N, tk.S))
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)

        self.tree = ttk.Treeview(tree_frame, columns=('ID', 'Отображение'), show='headings')
        self.tree.heading('ID', text='ID')
        self.tree.heading('Отображение', text='[Фамилия И.О.] (Фамилии И.О.) - Департамент')

        self.tree.column('ID', width=50)
        self.tree.column('Отображение', width=1000)

        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))

        self.refresh_list()

    def refresh_list(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        directors = self.db.get_all_persons(role='director')

        for director in directors:
            dir_id = director['id']
            last_name = director['last_name']
            first_name = director['first_name']
            middle_name = director['middle_name']
            department = director.get('department', '')

            nominative = f"{last_name} {first_name}.{middle_name}."

            declined_datv = self.decliner.decline_full_name(
                last_name, first_name, middle_name, 'datv', person_id=dir_id
            )
            first_init_datv = declined_datv['first_name'][0] if declined_datv['first_name'] else first_name[0]
            middle_init_datv = declined_datv['middle_name'][0] if declined_datv['middle_name'] else middle_name[0]
            dative = f"{declined_datv['last_name']} {first_init_datv}.{middle_init_datv}."

            display = f"[{nominative}] ({dative}) - {department}"

            self.tree.insert('', tk.END, values=(dir_id, display), tags=(last_name, first_name, middle_name, department))

    def add_director(self):
        dialog = tk.Toplevel(self.window)
        dialog.title("Добавить директора")
        dialog.geometry("600x300")

        ttk.Label(dialog, text="Фамилия:").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        last_name_entry = ttk.Entry(dialog, width=50)
        last_name_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Имя:").grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        first_name_entry = ttk.Entry(dialog, width=50)
        first_name_entry.grid(row=1, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Отчество:").grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)
        middle_name_entry = ttk.Entry(dialog, width=50)
        middle_name_entry.grid(row=2, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Департамент:").grid(row=3, column=0, padx=10, pady=10, sticky=tk.W)
        department_entry = ttk.Entry(dialog, width=50)
        department_entry.grid(row=3, column=1, padx=10, pady=10)

        def save():
            last_name = last_name_entry.get().strip()
            first_name = first_name_entry.get().strip()
            middle_name = middle_name_entry.get().strip()
            department = department_entry.get().strip()

            if not all([last_name, first_name, middle_name, department]):
                messagebox.showwarning("Предупреждение", "Заполните все поля")
                return

            if self.db.add_director(last_name, first_name, middle_name, department):
                messagebox.showinfo("Успех", "Директор добавлен")
                self.refresh_list()
                dialog.destroy()
            else:
                messagebox.showerror("Ошибка", "Не удалось добавить директора (возможно, уже существует)")

        ttk.Button(dialog, text="Сохранить", command=save).grid(row=4, column=0, columnspan=2, pady=20)

    def edit_director(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите директора для редактирования")
            return

        item = self.tree.item(selected[0])
        dir_id = item['values'][0]
        last_name, first_name, middle_name, department = item['tags']

        dialog = tk.Toplevel(self.window)
        dialog.title("Редактировать директора")
        dialog.geometry("600x300")

        ttk.Label(dialog, text="Фамилия:").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        last_name_entry = ttk.Entry(dialog, width=50)
        last_name_entry.insert(0, last_name)
        last_name_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Имя:").grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        first_name_entry = ttk.Entry(dialog, width=50)
        first_name_entry.insert(0, first_name)
        first_name_entry.grid(row=1, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Отчество:").grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)
        middle_name_entry = ttk.Entry(dialog, width=50)
        middle_name_entry.insert(0, middle_name)
        middle_name_entry.grid(row=2, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Департамент:").grid(row=3, column=0, padx=10, pady=10, sticky=tk.W)
        department_entry = ttk.Entry(dialog, width=50)
        department_entry.insert(0, department)
        department_entry.grid(row=3, column=1, padx=10, pady=10)

        def save():
            new_last_name = last_name_entry.get().strip()
            new_first_name = first_name_entry.get().strip()
            new_middle_name = middle_name_entry.get().strip()
            new_department = department_entry.get().strip()

            if not all([new_last_name, new_first_name, new_middle_name, new_department]):
                messagebox.showwarning("Предупреждение", "Заполните все поля")
                return

            if self.db.update_director(dir_id, new_last_name, new_first_name, new_middle_name, new_department):
                messagebox.showinfo("Успех", "Директор обновлен")
                self.refresh_list()
                dialog.destroy()
            else:
                messagebox.showerror("Ошибка", "Не удалось обновить директора")

        ttk.Button(dialog, text="Сохранить", command=save).grid(row=4, column=0, columnspan=2, pady=20)

    def delete_director(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите директора для удаления")
            return

        item = self.tree.item(selected[0])
        dir_id = item['values'][0]

        if messagebox.askyesno("Подтверждение", "Вы уверены, что хотите удалить этого директора?"):
            if self.db.delete_director(dir_id):
                messagebox.showinfo("Успех", "Директор удален")
                self.refresh_list()
            else:
                messagebox.showerror("Ошибка", "Не удалось удалить директора")

    def edit_declensions(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите директора для редактирования склонений")
            return

        item = self.tree.item(selected[0])
        person_id = item['values'][0]
        
        persons = self.db.get_all_persons(role='director')
        person = next((p for p in persons if p['id'] == person_id), None)
        if not person:
            messagebox.showerror("Ошибка", "Директор не найден")
            return
        
        last_name = person['last_name']
        first_name = person['first_name']
        middle_name = person['middle_name']
        department = person.get('department', '')

        gender = self.decliner.detect_gender(first_name, middle_name)
        full_name = f"{last_name} {first_name} {middle_name}"

        manual_decls = self.db.get_all_manual_declensions(person_id)

        auto_nomn = self.decliner.decline_full_name(last_name, first_name, middle_name, 'nomn')
        auto_gent = self.decliner.decline_full_name(last_name, first_name, middle_name, 'gent')
        auto_datv = self.decliner.decline_full_name(last_name, first_name, middle_name, 'datv')
        auto_accs = self.decliner.decline_full_name(last_name, first_name, middle_name, 'accs')

        dialog = tk.Toplevel(self.window)
        dialog.title(f"Редактирование склонений: {full_name}")
        dialog.geometry("700x450")

        ttk.Label(dialog, text="Ручное редактирование склонений ФИО", 
                 font=('Arial', 12, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)

        ttk.Label(dialog, text="Падеж", font=('Arial', 10, 'bold')).grid(row=1, column=0, padx=10, pady=5)
        ttk.Label(dialog, text="Автоматическое склонение", font=('Arial', 10, 'bold')).grid(row=1, column=1, padx=10, pady=5)
        ttk.Label(dialog, text="Ручное склонение (редактируемое)", font=('Arial', 10, 'bold')).grid(row=1, column=2, padx=10, pady=5)

        entries = {}
        cases = [
            ('nomn', 'Именительный (кто?)', auto_nomn),
            ('gent', 'Родительный (кого?)', auto_gent),
            ('datv', 'Дательный (кому?)', auto_datv),
            ('accs', 'Винительный (кого?)', auto_accs)
        ]

        for idx, (case_code, case_name, auto_value) in enumerate(cases, start=2):
            ttk.Label(dialog, text=case_name).grid(row=idx, column=0, padx=10, pady=10, sticky=tk.W)
            
            ttk.Label(dialog, text=auto_value, relief=tk.SUNKEN, 
                     background='#f0f0f0', width=30).grid(row=idx, column=1, padx=10, pady=10)
            
            entry = ttk.Entry(dialog, width=30)
            current_value = manual_decls.get(case_code, auto_value)
            entry.insert(0, current_value)
            entry.grid(row=idx, column=2, padx=10, pady=10)
            entries[case_code] = entry

        def save_declensions():
            for case_code, entry in entries.items():
                custom_value = entry.get().strip()
                if custom_value:
                    self.db.set_manual_declension(person_id, case_code, custom_value)
            
            messagebox.showinfo("Успех", "Склонения сохранены")
            self.refresh_list()
            dialog.destroy()

        def reset_to_auto():
            if messagebox.askyesno("Подтверждение", "Удалить все ручные склонения и вернуться к автоматическим?"):
                self.db.delete_manual_declension(person_id)
                messagebox.showinfo("Успех", "Ручные склонения удалены")
                self.refresh_list()
                dialog.destroy()

        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=6, column=0, columnspan=3, pady=20)

        ttk.Button(button_frame, text="Сохранить", command=save_declensions).grid(row=0, column=0, padx=10)
        ttk.Button(button_frame, text="Сбросить к автоматическим", command=reset_to_auto).grid(row=0, column=1, padx=10)
        ttk.Button(button_frame, text="Отмена", command=dialog.destroy).grid(row=0, column=2, padx=10)


class InspectorChiefsManager:
    def __init__(self, window, db, decliner):
        self.window = window
        self.db = db
        self.decliner = decliner

        main_frame = ttk.Frame(window, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        window.columnconfigure(0, weight=1)
        window.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(1, weight=1)

        title = ttk.Label(main_frame, text="Справочник начальников инспекций", 
                         font=('Arial', 14, 'bold'))
        title.grid(row=0, column=0, pady=10)

        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, pady=5, sticky=(tk.W, tk.E))

        ttk.Button(button_frame, text="Добавить", command=self.add_chief).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Редактировать", command=self.edit_chief).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="Редактировать склонения", command=self.edit_declensions).grid(row=0, column=2, padx=5)
        ttk.Button(button_frame, text="Удалить", command=self.delete_chief).grid(row=0, column=3, padx=5)
        ttk.Button(button_frame, text="Обновить", command=self.refresh_list).grid(row=0, column=4, padx=5)

        tree_frame = ttk.Frame(main_frame)
        tree_frame.grid(row=2, column=0, pady=10, sticky=(tk.W, tk.E, tk.N, tk.S))
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)

        self.tree = ttk.Treeview(tree_frame, columns=('ID', 'Отображение'), show='headings')
        self.tree.heading('ID', text='ID')
        self.tree.heading('Отображение', text='[Фамилия И.О.] (Фамилии И.О.)')

        self.tree.column('ID', width=50)
        self.tree.column('Отображение', width=800)

        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))

        self.refresh_list()

    def refresh_list(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        chiefs = self.db.get_all_persons(role='inspector')

        for chief in chiefs:
            chief_id = chief['id']
            last_name_nom = chief['last_name']
            first_init = chief['first_name']
            middle_init = chief['middle_name']

            nominative = f"{last_name_nom} {first_init}.{middle_init}."

            declined_datv = self.decliner.decline_full_name(
                last_name_nom, first_init, middle_init, 'datv', person_id=chief_id
            )
            first_init_datv = declined_datv['first_name'][0] if declined_datv['first_name'] else first_init[0]
            middle_init_datv = declined_datv['middle_name'][0] if declined_datv['middle_name'] else middle_init[0]
            dative = f"{declined_datv['last_name']} {first_init_datv}.{middle_init_datv}."

            display = f"[{nominative}] ({dative})"

            self.tree.insert('', tk.END, values=(chief_id, display), 
                           tags=(last_name_nom, declined_datv['last_name'], first_init, middle_init))

    def add_chief(self):
        dialog = tk.Toplevel(self.window)
        dialog.title("Добавить начальника инспекции")
        dialog.geometry("700x400")

        ttk.Label(dialog, text="Фамилия (именительный падеж):").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        last_name_nom_entry = ttk.Entry(dialog, width=50)
        last_name_nom_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Фамилия (дательный падеж):").grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        last_name_dat_entry = ttk.Entry(dialog, width=50)
        last_name_dat_entry.grid(row=1, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Инициал имени (А, Б, В...):").grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)
        first_init_entry = ttk.Entry(dialog, width=50)
        first_init_entry.grid(row=2, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Инициал отчества (А, Б, В...):").grid(row=3, column=0, padx=10, pady=10, sticky=tk.W)
        middle_init_entry = ttk.Entry(dialog, width=50)
        middle_init_entry.grid(row=3, column=1, padx=10, pady=10)

        def auto_fill_dative():
            nom = last_name_nom_entry.get().strip()
            first = first_init_entry.get().strip()
            middle = middle_init_entry.get().strip()

            if nom and first and middle:
                gender = 'male'
                if first in ['Е', 'И', 'Ю', 'Я', 'А']:
                    if first in ['Е', 'И', 'Ю', 'Я']:
                        gender = 'female'

                dative = self.decliner.decline_name_with_gender(nom, 'datv', gender, 'Surn')
                last_name_dat_entry.delete(0, tk.END)
                last_name_dat_entry.insert(0, dative)

        ttk.Button(dialog, text="Автозаполнить дательный падеж", command=auto_fill_dative).grid(row=4, column=0, columnspan=2, pady=10)

        def save():
            last_name_nom = last_name_nom_entry.get().strip()
            last_name_dat = last_name_dat_entry.get().strip()
            first_init = first_init_entry.get().strip().upper()
            middle_init = middle_init_entry.get().strip().upper()

            if not all([last_name_nom, last_name_dat, first_init, middle_init]):
                messagebox.showwarning("Предупреждение", "Заполните все поля")
                return

            if len(first_init) != 1 or len(middle_init) != 1:
                messagebox.showwarning("Предупреждение", "Инициалы должны быть одной буквой")
                return

            if self.db.add_inspector_chief(last_name_nom, last_name_dat, first_init, middle_init, ''):
                messagebox.showinfo("Успех", "Начальник инспекции добавлен")
                self.refresh_list()
                dialog.destroy()
            else:
                messagebox.showerror("Ошибка", "Не удалось добавить начальника инспекции")

        ttk.Button(dialog, text="Сохранить", command=save).grid(row=5, column=0, columnspan=2, pady=20)

    def edit_chief(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите начальника для редактирования")
            return

        item = self.tree.item(selected[0])
        chief_id = item['values'][0]
        last_name_nom, last_name_dat, first_init, middle_init = item['tags']

        dialog = tk.Toplevel(self.window)
        dialog.title("Редактировать начальника инспекции")
        dialog.geometry("700x350")

        ttk.Label(dialog, text="Фамилия (именительный падеж):").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        last_name_nom_entry = ttk.Entry(dialog, width=50)
        last_name_nom_entry.insert(0, last_name_nom)
        last_name_nom_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Фамилия (дательный падеж):").grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        last_name_dat_entry = ttk.Entry(dialog, width=50)
        last_name_dat_entry.insert(0, last_name_dat)
        last_name_dat_entry.grid(row=1, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Инициал имени (А, Б, В...):").grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)
        first_init_entry = ttk.Entry(dialog, width=50)
        first_init_entry.insert(0, first_init)
        first_init_entry.grid(row=2, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Инициал отчества (А, Б, В...):").grid(row=3, column=0, padx=10, pady=10, sticky=tk.W)
        middle_init_entry = ttk.Entry(dialog, width=50)
        middle_init_entry.insert(0, middle_init)
        middle_init_entry.grid(row=3, column=1, padx=10, pady=10)

        def auto_fill_dative():
            nom = last_name_nom_entry.get().strip()
            first = first_init_entry.get().strip()
            middle = middle_init_entry.get().strip()

            if nom and first and middle:
                gender = 'male'
                if first in ['Е', 'И', 'Ю', 'Я', 'А']:
                    if first in ['Е', 'И', 'Ю', 'Я']:
                        gender = 'female'

                dative = self.decliner.decline_name_with_gender(nom, 'datv', gender, 'Surn')
                last_name_dat_entry.delete(0, tk.END)
                last_name_dat_entry.insert(0, dative)

        ttk.Button(dialog, text="Автозаполнить дательный падеж", command=auto_fill_dative).grid(row=4, column=0, columnspan=2, pady=10)

        def save():
            new_last_name_nom = last_name_nom_entry.get().strip()
            new_last_name_dat = last_name_dat_entry.get().strip()
            new_first_init = first_init_entry.get().strip().upper()
            new_middle_init = middle_init_entry.get().strip().upper()

            if not all([new_last_name_nom, new_last_name_dat, new_first_init, new_middle_init]):
                messagebox.showwarning("Предупреждение", "Заполните все поля")
                return

            if len(new_first_init) != 1 or len(new_middle_init) != 1:
                messagebox.showwarning("Предупреждение", "Инициалы должны быть одной буквой")
                return

            if self.db.update_inspector_chief(chief_id, new_last_name_nom, new_last_name_dat, 
                                             new_first_init, new_middle_init, ''):
                messagebox.showinfo("Успех", "Начальник инспекции обновлен")
                self.refresh_list()
                dialog.destroy()
            else:
                messagebox.showerror("Ошибка", "Не удалось обновить начальника инспекции")

        ttk.Button(dialog, text="Сохранить", command=save).grid(row=5, column=0, columnspan=2, pady=20)

    def delete_chief(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите начальника для удаления")
            return

        item = self.tree.item(selected[0])
        chief_id = item['values'][0]

        if messagebox.askyesno("Подтверждение", "Вы уверены, что хотите удалить этого начальника инспекции?"):
            if self.db.delete_inspector_chief(chief_id):
                messagebox.showinfo("Успех", "Начальник инспекции удален")
                self.refresh_list()
            else:
                messagebox.showerror("Ошибка", "Не удалось удалить начальника инспекции")

    def edit_declensions(self):
        selected = self.tree.selection()
        if not selected:
            messagebox.showwarning("Предупреждение", "Выберите начальника для редактирования склонений")
            return

        item = self.tree.item(selected[0])
        person_id = item['values'][0]
        
        persons = self.db.get_all_persons(role='inspector')
        person = next((p for p in persons if p['id'] == person_id), None)
        if not person:
            messagebox.showerror("Ошибка", "Начальник инспекции не найден")
            return
        
        last_name = person['last_name']
        first_name = person['first_name']
        middle_name = person['middle_name']

        gender = self.decliner.detect_gender(first_name, middle_name)
        full_name = f"{last_name} {first_name} {middle_name}"

        manual_decls = self.db.get_all_manual_declensions(person_id)

        auto_nomn = self.decliner.decline_full_name(last_name, first_name, middle_name, 'nomn')
        auto_gent = self.decliner.decline_full_name(last_name, first_name, middle_name, 'gent')
        auto_datv = self.decliner.decline_full_name(last_name, first_name, middle_name, 'datv')
        auto_accs = self.decliner.decline_full_name(last_name, first_name, middle_name, 'accs')

        dialog = tk.Toplevel(self.window)
        dialog.title(f"Редактирование склонений: {full_name}")
        dialog.geometry("700x450")

        ttk.Label(dialog, text="Ручное редактирование склонений ФИО", 
                 font=('Arial', 12, 'bold')).grid(row=0, column=0, columnspan=3, pady=10)

        ttk.Label(dialog, text="Падеж", font=('Arial', 10, 'bold')).grid(row=1, column=0, padx=10, pady=5)
        ttk.Label(dialog, text="Автоматическое склонение", font=('Arial', 10, 'bold')).grid(row=1, column=1, padx=10, pady=5)
        ttk.Label(dialog, text="Ручное склонение (редактируемое)", font=('Arial', 10, 'bold')).grid(row=1, column=2, padx=10, pady=5)

        entries = {}
        cases = [
            ('nomn', 'Именительный (кто?)', auto_nomn),
            ('gent', 'Родительный (кого?)', auto_gent),
            ('datv', 'Дательный (кому?)', auto_datv),
            ('accs', 'Винительный (кого?)', auto_accs)
        ]

        for idx, (case_code, case_name, auto_value) in enumerate(cases, start=2):
            ttk.Label(dialog, text=case_name).grid(row=idx, column=0, padx=10, pady=10, sticky=tk.W)
            
            ttk.Label(dialog, text=auto_value, relief=tk.SUNKEN, 
                     background='#f0f0f0', width=30).grid(row=idx, column=1, padx=10, pady=10)
            
            entry = ttk.Entry(dialog, width=30)
            current_value = manual_decls.get(case_code, auto_value)
            entry.insert(0, current_value)
            entry.grid(row=idx, column=2, padx=10, pady=10)
            entries[case_code] = entry

        def save_declensions():
            for case_code, entry in entries.items():
                custom_value = entry.get().strip()
                if custom_value:
                    self.db.set_manual_declension(person_id, case_code, custom_value)
            
            messagebox.showinfo("Успех", "Склонения сохранены")
            self.refresh_list()
            dialog.destroy()

        def reset_to_auto():
            if messagebox.askyesno("Подтверждение", "Удалить все ручные склонения и вернуться к автоматическим?"):
                self.db.delete_manual_declension(person_id)
                messagebox.showinfo("Успех", "Ручные склонения удалены")
                self.refresh_list()
                dialog.destroy()

        button_frame = ttk.Frame(dialog)
        button_frame.grid(row=6, column=0, columnspan=3, pady=20)

        ttk.Button(button_frame, text="Сохранить", command=save_declensions).grid(row=0, column=0, padx=10)
        ttk.Button(button_frame, text="Сбросить к автоматическим", command=reset_to_auto).grid(row=0, column=1, padx=10)
        ttk.Button(button_frame, text="Отмена", command=dialog.destroy).grid(row=0, column=2, padx=10)


class OrganizationsManager:
    def __init__(self, window, db, root):
        self.window = window
        self.db = db
        self.root = root

        main_frame = ttk.Frame(window, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        window.columnconfigure(0, weight=1)
        window.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(2, weight=1)

        title = ttk.Label(main_frame, text="Справочник организаций", 
                         font=('Arial', 14, 'bold'))
        title.grid(row=0, column=0, pady=10)

        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=1, column=0, pady=5, sticky=(tk.W, tk.E))

        ttk.Button(button_frame, text="Добавить", command=self.add_organization).grid(row=0, column=0, padx=5)
        ttk.Button(button_frame, text="Редактировать", command=self.edit_organization).grid(row=0, column=1, padx=5)
        ttk.Button(button_frame, text="Удалить", command=self.delete_organization).grid(row=0, column=2, padx=5)
        ttk.Button(button_frame, text="Обновить", command=self.refresh_list).grid(row=0, column=3, padx=5)
        
        sync_btn = ttk.Button(button_frame, text="🔄 Синхронизация с data.mos.ru", command=self.open_sync_dialog)
        sync_btn.grid(row=0, column=4, padx=15)

        tree_frame = ttk.Frame(main_frame)
        tree_frame.grid(row=2, column=0, pady=10, sticky=(tk.W, tk.E, tk.N, tk.S))
        tree_frame.columnconfigure(0, weight=1)
        tree_frame.rowconfigure(0, weight=1)

        self.tree = ttk.Treeview(tree_frame, columns=('ID', 'Название', 'Краткое', 'ИНН', 'ОГРН'), show='headings')
        self.tree.heading('ID', text='ID')
        self.tree.heading('Название', text='Название организации')
        self.tree.heading('Краткое', text='Краткое название')
        self.tree.heading('ИНН', text='ИНН')
        self.tree.heading('ОГРН', text='ОГРН')

        self.tree.column('ID', width=50)
        self.tree.column('Название', width=400)
        self.tree.column('Краткое', width=200)
        self.tree.column('ИНН', width=120)
        self.tree.column('ОГРН', width=150)

        scrollbar = ttk.Scrollbar(tree_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)

        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))

        self.refresh_list()

    def refresh_list(self):
        for item in self.tree.get_children():
            self.tree.delete(item)

        departments = self.db.get_all_departments(active_only=True)

        for dept in departments:
            dept_id = dept['id']
            name = dept['name']
            short_name = dept.get('short_name', '')
            inn = dept.get('inn', '')
            ogrn = dept.get('ogrn', '')

            self.tree.insert('', tk.END, values=(dept_id, name, short_name, inn, ogrn))

    def add_organization(self):
        dialog = tk.Toplevel(self.window)
        dialog.title("Добавить организацию")
        dialog.geometry("600x300")
        dialog.transient(self.window)
        dialog.grab_set()

        ttk.Label(dialog, text="Название:").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        name_entry = ttk.Entry(dialog, width=50)
        name_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Краткое название:").grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        short_name_entry = ttk.Entry(dialog, width=50)
        short_name_entry.grid(row=1, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="ИНН:").grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)
        inn_entry = ttk.Entry(dialog, width=50)
        inn_entry.grid(row=2, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="ОГРН:").grid(row=3, column=0, padx=10, pady=10, sticky=tk.W)
        ogrn_entry = ttk.Entry(dialog, width=50)
        ogrn_entry.grid(row=3, column=1, padx=10, pady=10)

        def save():
            name = name_entry.get().strip()
            if not name:
                tk.messagebox.showerror("Ошибка", "Название обязательно")
                return

            short_name = short_name_entry.get().strip()
            inn = inn_entry.get().strip()
            ogrn = ogrn_entry.get().strip()

            if self.db.add_department(name, short_name, inn, ogrn, ""):
                tk.messagebox.showinfo("Успех", "Организация добавлена")
                self.refresh_list()
                dialog.destroy()
            else:
                tk.messagebox.showerror("Ошибка", "Не удалось добавить организацию (возможно, уже существует)")

        ttk.Button(dialog, text="Сохранить", command=save).grid(row=4, column=0, pady=20, padx=10)
        ttk.Button(dialog, text="Отмена", command=dialog.destroy).grid(row=4, column=1, pady=20)

    def edit_organization(self):
        selection = self.tree.selection()
        if not selection:
            tk.messagebox.showwarning("Предупреждение", "Выберите организацию для редактирования")
            return

        item_values = self.tree.item(selection[0], 'values')
        dept_id = int(item_values[0])
        current_name = item_values[1]
        current_short = item_values[2]
        current_inn = item_values[3]
        current_ogrn = item_values[4]

        dialog = tk.Toplevel(self.window)
        dialog.title("Редактировать организацию")
        dialog.geometry("600x300")
        dialog.transient(self.window)
        dialog.grab_set()

        ttk.Label(dialog, text="Название:").grid(row=0, column=0, padx=10, pady=10, sticky=tk.W)
        name_entry = ttk.Entry(dialog, width=50)
        name_entry.insert(0, current_name)
        name_entry.grid(row=0, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="Краткое название:").grid(row=1, column=0, padx=10, pady=10, sticky=tk.W)
        short_name_entry = ttk.Entry(dialog, width=50)
        short_name_entry.insert(0, current_short)
        short_name_entry.grid(row=1, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="ИНН:").grid(row=2, column=0, padx=10, pady=10, sticky=tk.W)
        inn_entry = ttk.Entry(dialog, width=50)
        inn_entry.insert(0, current_inn)
        inn_entry.grid(row=2, column=1, padx=10, pady=10)

        ttk.Label(dialog, text="ОГРН:").grid(row=3, column=0, padx=10, pady=10, sticky=tk.W)
        ogrn_entry = ttk.Entry(dialog, width=50)
        ogrn_entry.insert(0, current_ogrn)
        ogrn_entry.grid(row=3, column=1, padx=10, pady=10)

        def save():
            name = name_entry.get().strip()
            if not name:
                tk.messagebox.showerror("Ошибка", "Название обязательно")
                return

            short_name = short_name_entry.get().strip()
            inn = inn_entry.get().strip()
            ogrn = ogrn_entry.get().strip()

            if self.db.update_department(dept_id, name, short_name, inn, ogrn, ""):
                tk.messagebox.showinfo("Успех", "Организация обновлена")
                self.refresh_list()
                dialog.destroy()
            else:
                tk.messagebox.showerror("Ошибка", "Не удалось обновить организацию")

        ttk.Button(dialog, text="Сохранить", command=save).grid(row=4, column=0, pady=20, padx=10)
        ttk.Button(dialog, text="Отмена", command=dialog.destroy).grid(row=4, column=1, pady=20)

    def delete_organization(self):
        selection = self.tree.selection()
        if not selection:
            tk.messagebox.showwarning("Предупреждение", "Выберите организацию для удаления")
            return

        item_values = self.tree.item(selection[0], 'values')
        dept_id = int(item_values[0])
        dept_name = item_values[1]

        if tk.messagebox.askyesno("Подтверждение", f"Удалить организацию '{dept_name}'?"):
            if self.db.delete_department(dept_id):
                tk.messagebox.showinfo("Успех", "Организация удалена")
                self.refresh_list()
            else:
                tk.messagebox.showerror("Ошибка", "Не удалось удалить организацию")

    def open_sync_dialog(self):
        sync_window = tk.Toplevel(self.window)
        sync_window.title("Синхронизация с data.mos.ru")
        sync_window.geometry("700x500")
        sync_window.transient(self.window)
        sync_window.grab_set()

        main_frame = ttk.Frame(sync_window, padding="15")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        sync_window.columnconfigure(0, weight=1)
        sync_window.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(3, weight=1)

        ttk.Label(main_frame, text="Синхронизация организаций с data.mos.ru", 
                 font=('Arial', 12, 'bold')).grid(row=0, column=0, pady=10)

        api_frame = ttk.Frame(main_frame)
        api_frame.grid(row=1, column=0, pady=10, sticky=(tk.W, tk.E))
        
        ttk.Label(api_frame, text="API ключ:").grid(row=0, column=0, padx=5, sticky=tk.W)
        api_key_entry = ttk.Entry(api_frame, width=60, show="*")
        api_key_entry.grid(row=0, column=1, padx=5)

        progress_frame = ttk.Frame(main_frame)
        progress_frame.grid(row=2, column=0, pady=10, sticky=(tk.W, tk.E))
        progress_frame.columnconfigure(0, weight=1)

        progress_label = ttk.Label(progress_frame, text="Готов к синхронизации")
        progress_label.grid(row=0, column=0, sticky=tk.W, pady=5)

        progress_bar = ttk.Progressbar(progress_frame, mode='determinate', maximum=100)
        progress_bar.grid(row=1, column=0, sticky=(tk.W, tk.E))

        log_frame = ttk.LabelFrame(main_frame, text="Лог синхронизации")
        log_frame.grid(row=3, column=0, pady=10, sticky=(tk.W, tk.E, tk.N, tk.S))
        log_frame.columnconfigure(0, weight=1)
        log_frame.rowconfigure(0, weight=1)

        log_text = tk.Text(log_frame, height=15, wrap=tk.WORD, state=tk.DISABLED)
        log_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        log_scroll = ttk.Scrollbar(log_frame, orient=tk.VERTICAL, command=log_text.yview)
        log_text.configure(yscrollcommand=log_scroll.set)
        log_scroll.grid(row=0, column=1, sticky=(tk.N, tk.S))

        def log_message(message):
            log_text.config(state=tk.NORMAL)
            log_text.insert(tk.END, message + "\n")
            log_text.see(tk.END)
            log_text.config(state=tk.DISABLED)

        def update_progress(current, total):
            if total > 0:
                percent = int((current / total) * 100)
                progress_bar['value'] = percent
                progress_label.config(text=f"Загружено: {current}/{total} записей ({percent}%)")

        def run_sync():
            api_key = api_key_entry.get().strip()
            if not api_key:
                tk.messagebox.showerror("Ошибка", "Введите API ключ")
                return

            sync_btn.config(state=tk.DISABLED)
            log_text.config(state=tk.NORMAL)
            log_text.delete('1.0', tk.END)
            log_text.config(state=tk.DISABLED)
            progress_bar['value'] = 0

            def sync_thread():
                try:
                    added, updated, skipped = sync_organizations_from_data_mos_ru(
                        self.db, api_key, 
                        progress_callback=lambda c, t: self.root.after(0, update_progress, c, t),
                        log_callback=lambda msg: self.root.after(0, log_message, msg)
                    )
                    
                    self.root.after(0, lambda: tk.messagebox.showinfo(
                        "Синхронизация завершена",
                        f"Добавлено: {added}\nОбновлено: {updated}\nПропущено: {skipped}"
                    ))
                    self.root.after(0, self.refresh_list)
                except Exception as e:
                    self.root.after(0, lambda: tk.messagebox.showerror("Ошибка", f"Ошибка синхронизации:\n{str(e)}"))
                finally:
                    self.root.after(0, lambda: sync_btn.config(state=tk.NORMAL))

            threading.Thread(target=sync_thread, daemon=True).start()

        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=4, column=0, pady=15)

        sync_btn = ttk.Button(button_frame, text="Начать синхронизацию", command=run_sync)
        sync_btn.grid(row=0, column=0, padx=10)

        ttk.Button(button_frame, text="Закрыть", command=sync_window.destroy).grid(row=0, column=1, padx=10)


if __name__ == "__main__":
    # Print diagnostic information on startup
    from src.version_info import get_full_version_string
    version_string = get_full_version_string()
    
    print("="*70)
    print(f"ЗАПУСК ПРИЛОЖЕНИЯ PDF Parser ОАТИ {version_string}")
    print("="*70)
    print(f"Текущая рабочая директория: {os.getcwd()}")
    print(f"Абсолютный путь: {os.path.abspath('.')}")
    print(f"Папка output существует: {os.path.exists('output')}")
    print(f"Папка templates существует: {os.path.exists('templates')}")
    print(f"Папка database существует: {os.path.exists('database')}")

    # Ensure output directory exists
    output_dir = Path("output")
    if not output_dir.exists():
        output_dir.mkdir(parents=True)
        print(f"[CREATE] Создана папка: {output_dir.absolute()}")
    else:
        print(f"[OK] Папка output: {output_dir.absolute()}")

    print("="*70)

    setup_application()

    try:
        root = tkdnd.TkinterDnD.Tk()
    except:
        root = tk.Tk()

    templates_dir = Path("templates")
    template_count = len(list(templates_dir.glob("*.docx"))) if templates_dir.exists() else 0

    if template_count == 0:
        messagebox.showwarning(
            "Шаблоны не найдены",
            f"Шаблоны Word документов не обнаружены!\n\n"
            f"Для работы приложения необходимо добавить файлы .docx\n"
            f"в папку: {templates_dir.absolute()}\n\n"
            f"Приложение откроется, но генерация документов будет недоступна."
        )

    app = PDFProcessorApp(root)
    root.mainloop()