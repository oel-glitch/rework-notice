"""
Word document generator module for PDF Parser application.

This module handles the generation of Word documents from templates using
configuration-based template selection and citizen data.
"""

import logging
import os
import re
import warnings
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple, Union

from docx import Document
from docx.text.paragraph import Paragraph

from src.config_loader import config_loader
from src.recipient_types import (
    classify_recipients,
    RecipientInfo,
    RecipientsClassification,
    RecipientRole
)
from src.gost_normalizer import GOSTNormalizer


logger = logging.getLogger(__name__)


class WordGenerator:
    """
    Generates Word documents from templates with dynamic data replacement.
    
    This class provides functionality for selecting appropriate templates based on
    configuration rules, replacing placeholders with actual data, and generating
    properly formatted output documents.
    
    Attributes:
        templates_dir: Directory path containing document templates.
        output_dir: Directory path for generated output documents.
        template_config: Configuration for template selection and placeholders.
    """
    
    def __init__(self, templates_dir: str = "templates", output_dir: str = "output") -> None:
        """
        Initialize the Word document generator.
        
        Args:
            templates_dir: Path to directory containing template files. Defaults to "templates".
            output_dir: Path to directory for output documents. Defaults to "output".
            
        Raises:
            ValueError: If configuration cannot be loaded.
        """
        self.templates_dir: str = templates_dir
        self.output_dir: str = output_dir
        
        # Template cache for performance optimization during batch processing
        # Cache stores (file_mtime, document_bytes) to detect template changes
        self._template_cache: Dict[str, Tuple[float, bytes]] = {}
        logger.debug("Template cache initialized for batch processing optimization")
        
        # Create output directory with detailed logging
        output_path = Path(output_dir)
        if not output_path.exists():
            output_path.mkdir(parents=True, exist_ok=True)
            logger.info(f"Created output directory: {output_path.absolute()}")
        else:
            logger.info(f"Output directory exists: {output_path.absolute()}")
        
        # Verify write permissions
        test_file = output_path / ".write_test"
        try:
            test_file.touch()
            test_file.unlink()
            logger.info(f"Output directory is writable: {output_path.absolute()}")
        except Exception as e:
            logger.error(f"Output directory is NOT writable: {output_path.absolute()} - {e}")
            raise PermissionError(f"Cannot write to output directory: {output_path.absolute()}")
        
        logger.info(f"WordGenerator initialized with templates_dir='{templates_dir}', output_dir='{output_dir}'")
        
        # Initialize GOST normalizer for optional full normalization
        # NOTE: Currently NOT used by default to preserve backward compatibility
        # Use apply_full_gost_normalization() for explicit normalization
        self.gost_normalizer = GOSTNormalizer()
        logger.debug("GOST normalizer initialized (optional)")
        
        try:
            self.template_config: Dict[str, Any] = config_loader.load('templates_mapping')
            logger.debug("Template configuration loaded successfully")
        except Exception as e:
            logger.error(f"Failed to load template configuration: {e}")
            raise ValueError(f"Cannot initialize WordGenerator: configuration error - {e}")
    
    def _evaluate_template_condition(
        self, 
        condition: str, 
        departments: List[str], 
        has_inspector: bool
    ) -> bool:
        """
        Evaluate whether a template condition is satisfied.
        
        Args:
            condition: Condition name from configuration.
            departments: List of department names.
            has_inspector: Whether an inspector is involved.
            
        Returns:
            True if condition is satisfied, False otherwise.
        """
        num_departments = len(departments)
        has_prefecture = any('префектура' in dept.lower() for dept in departments)
        
        if condition == "single_department_without_inspector":
            return num_departments == 1 and not has_inspector
        elif condition == "prefecture_recipient":
            return has_prefecture or has_inspector
        elif condition == "multiple_departments_or_inspector":
            return num_departments > 1 or has_inspector
        
        logger.warning(f"Unknown template condition: {condition}")
        return False
    
    def _find_template_file(self, pattern: str) -> Optional[str]:
        """
        Find template file matching the given pattern.
        
        Args:
            pattern: Glob pattern for template filename.
            
        Returns:
            Full path to template file if found, None otherwise.
        """
        try:
            templates_path = Path(self.templates_dir)
            matching_files = list(templates_path.glob(pattern))
            
            if matching_files:
                template_file = str(matching_files[0])
                logger.debug(f"Found template file: {template_file} for pattern: {pattern}")
                return template_file
            else:
                logger.warning(f"No template file found matching pattern: {pattern}")
                return None
        except Exception as e:
            logger.error(f"Error searching for template with pattern '{pattern}': {e}")
            return None
    
    def select_template_by_recipients(
        self,
        recipients: RecipientsClassification
    ) -> Tuple[str, bool, str]:
        """
        Select template based on typed recipient classification.
        
        Uses the new template_selection_rules matrix from configuration to determine
        which template should be used based on the combination of recipient types.
        
        Args:
            recipients: Classified recipient information.
            
        Returns:
            Tuple of (template_path, union_included, union_suffix).
            
        Raises:
            FileNotFoundError: If no suitable template file can be found.
        """
        combination_key = recipients.get_combination_key()
        logger.info(f"Selecting template by recipients: combination_key='{combination_key}', {recipients}")
        
        try:
            selection_rules = self.template_config.get('template_selection_rules', {})
            
            if combination_key in selection_rules:
                rule = selection_rules[combination_key]
                pattern = rule.get('template_pattern', 'template_ch3*.docx')
                union_included = rule.get('union_included', False)
                union_suffix = rule.get('union_suffix', '')
                
                template_path = self._find_template_file(pattern)
                
                if template_path:
                    logger.info(
                        f"Selected template using new rules: {template_path} "
                        f"(combination: {combination_key}, union_included: {union_included})"
                    )
                    return template_path, union_included, union_suffix
                else:
                    logger.warning(f"Template pattern '{pattern}' not found for combination '{combination_key}'")
            else:
                # Emit deprecation warning ONLY when combination_key is missing
                logger.warning(
                    f"No template selection rule found for combination '{combination_key}', "
                    f"falling back to deprecated logic"
                )
                warnings.warn(
                    f"Template selection rule for combination '{combination_key}' not found. "
                    f"Falling back to deprecated_law_parts configuration. "
                    f"Please update templates_mapping.json to include this combination.",
                    DeprecationWarning,
                    stacklevel=2
                )
            
            # Fallback to deprecated_law_parts
            
            deprecated_config = self.template_config.get('deprecated_law_parts', {})
            
            if recipients.has_prefecture:
                config_key = 'ch4_prefecture'
            elif recipients.has_multiple_recipients or recipients.has_oati:
                config_key = 'ch4_multiple'
            else:
                config_key = 'ch3'
            
            deprecated_rule = deprecated_config.get(config_key, {})
            pattern = deprecated_rule.get('template_pattern', 'template_ch3*.docx')
            union_included = deprecated_rule.get('union_included', False)
            union_suffix = deprecated_rule.get('union_suffix', '')
            
            template_path = self._find_template_file(pattern)
            
            if template_path:
                logger.info(
                    f"Selected template using deprecated fallback: {template_path} "
                    f"(config_key: {config_key}, union_included: {union_included})"
                )
                return template_path, union_included, union_suffix
            
            # Ultimate fallback - emit deprecation warning
            fallback_path = os.path.join(self.templates_dir, "template_ch3.docx")
            logger.error(f"No template found, attempting ultimate fallback: {fallback_path}")
            
            warnings.warn(
                f"No template file found for combination '{combination_key}' or deprecated config. "
                f"Using ultimate fallback: {fallback_path}. "
                f"This indicates a configuration or file system issue.",
                DeprecationWarning,
                stacklevel=2
            )
            
            if not os.path.exists(fallback_path):
                raise FileNotFoundError(f"Template file not found: {fallback_path}")
            
            return fallback_path, False, ''
            
        except Exception as e:
            logger.error(f"Error during template selection by recipients: {e}")
            raise
    
    def select_template(
        self, 
        law_part: str, 
        departments: List[str], 
        has_inspector: bool = False,
        recipients_classification: Optional[RecipientsClassification] = None,
        union_response: bool = False
    ) -> Union[str, Tuple[str, bool, str]]:
        """
        Select appropriate template based on law part and configuration rules.
        
        This method uses configuration-based template selection logic to determine
        which template file should be used based on the law part, number of departments,
        and presence of an inspector.
        
        Args:
            law_part: Law part identifier (e.g., "3" or "4").
            departments: List of department names involved.
            has_inspector: Whether an inspector is involved in the case.
            recipients_classification: Optional typed recipient classification.
                If provided, uses the new template selection logic and returns tuple.
            union_response: Whether ОАТИ will handle the case itself (True = Мухтарова variant).
            
        Returns:
            When recipients_classification is provided: Tuple of (template_path, union_included, union_suffix).
            When recipients_classification is None: str (template_path only) for backward compatibility.
            
        Raises:
            FileNotFoundError: If no suitable template file can be found.
        """
        if recipients_classification is not None:
            logger.info(f"Using new typed recipient classification for template selection (union_response={union_response})")
            
            # CRITICAL LOGIC: Override combination key if union_response is True
            # If ОАТИ handles the case (union_response=True), force the template to use
            # the special variant with additional paragraph (Мухтарова example)
            if union_response:
                logger.info("union_response=True: Forcing template_ch4_oati_handles variant (Мухтарова)")
                # Override combination to force OATI handling template
                original_key = recipients_classification.get_combination_key()
                
                # Force combination key to include "with_oati" suffix
                dept_count = len(recipients_classification.departments)
                if recipients_classification.has_prefecture:
                    forced_key = "prefecture_with_oati"
                elif dept_count == 0:
                    forced_key = "oati"
                elif dept_count == 1:
                    forced_key = "single_department_with_oati"
                else:
                    forced_key = "multiple_departments_with_oati"
                
                logger.info(f"Original combination: '{original_key}' → Forced combination: '{forced_key}'")
                
                # Manually select template for OATI handling cases
                selection_rules = self.template_config.get('template_selection_rules', {})
                if forced_key in selection_rules:
                    rule = selection_rules[forced_key]
                    pattern = rule.get('template_pattern', 'template_ch4_oati_handles*.docx')
                    union_included = rule.get('union_included', True)
                    union_suffix = rule.get('union_suffix', ' (далее – Объединение)')
                    
                    template_path = self._find_template_file(pattern)
                    if template_path:
                        logger.info(f"Selected template for union_response: {template_path}")
                        return (template_path, union_included, union_suffix)
            
            # Normal flow: use recipient classification as-is
            template_path, union_included, union_suffix = self.select_template_by_recipients(recipients_classification)
            logger.debug(f"Returning tuple: template_path='{template_path}', union_included={union_included}, union_suffix='{union_suffix}'")
            return (template_path, union_included, union_suffix)
        
        warnings.warn(
            "Using deprecated template selection logic based on law_part. "
            "Please use recipients_classification parameter for better type safety.",
            DeprecationWarning,
            stacklevel=2
        )
        logger.warning(
            "DEPRECATED: select_template() called without recipients_classification. "
            "Consider updating to use typed recipient classification."
        )
        
        logger.info(f"Selecting template for law_part='{law_part}', departments={len(departments)}, has_inspector={has_inspector}")
        
        try:
            law_parts_config = self.template_config.get('law_parts', {})
            
            if law_part == "3":
                ch3_config = law_parts_config.get('ch3', {})
                pattern = ch3_config.get('template_pattern', 'template_ch3.docx')
                template_path = self._find_template_file(pattern)
                
                if template_path:
                    logger.info(f"Selected template: {template_path} (law part 3)")
                    return template_path
                    
            elif law_part == "4":
                for template_key in ['ch4_prefecture', 'ch4_multiple']:
                    template_cfg = law_parts_config.get(template_key, {})
                    condition = template_cfg.get('condition', '')
                    
                    if self._evaluate_template_condition(condition, departments, has_inspector):
                        pattern = template_cfg.get('template_pattern', 'template_ch4_multiple.docx')
                        template_path = self._find_template_file(pattern)
                        
                        if template_path:
                            logger.info(f"Selected template: {template_path} (law part 4, condition: {condition})")
                            return template_path
            
            default_pattern = law_parts_config.get('ch3', {}).get('template_pattern', 'template_ch3.docx')
            default_path = self._find_template_file(default_pattern)
            
            if default_path:
                logger.warning(f"Using default template: {default_path}")
                return default_path
            
            fallback_path = os.path.join(self.templates_dir, "template_ch3.docx")
            logger.error(f"No template found, attempting fallback: {fallback_path}")
            
            if not os.path.exists(fallback_path):
                raise FileNotFoundError(f"Template file not found: {fallback_path}")
                
            return fallback_path
            
        except Exception as e:
            logger.error(f"Error during template selection: {e}")
            raise
    
    def generate_filename(
        self, 
        last_name: str, 
        first_initial: str, 
        middle_initial: str, 
        oati_number: str
    ) -> str:
        """
        Generate output filename based on citizen data.
        
        Args:
            last_name: Citizen's last name.
            first_initial: First letter of citizen's first name.
            middle_initial: First letter of citizen's middle name.
            oati_number: OATI document number.
            
        Returns:
            Generated filename for output document.
        """
        try:
            number_match = re.search(r'01-21-П-(\d+)/(\d+)', oati_number)
            
            if number_match:
                central_digits = number_match.group(1)
                year_suffix = number_match.group(2)
                
                if '-' in central_digits:
                    suffix = central_digits.split('-')[1]
                    central_base = central_digits.split('-')[0]
                    filename = (
                        f"Уведомление_гражданина_{last_name}_{first_initial}_"
                        f"{middle_initial}_{central_base}_{year_suffix}-{suffix}.docx"
                    )
                else:
                    filename = (
                        f"Уведомление_гражданина_{last_name}_{first_initial}_"
                        f"{middle_initial}_{central_digits}_{year_suffix}.docx"
                    )
            else:
                filename = f"Уведомление_гражданина_{last_name}_{first_initial}_{middle_initial}.docx"
            
            logger.debug(f"Generated filename: {filename}")
            return filename
            
        except Exception as e:
            logger.error(f"Error generating filename: {e}")
            fallback_filename = f"Уведомление_гражданина_{last_name}.docx"
            logger.warning(f"Using fallback filename: {fallback_filename}")
            return fallback_filename
    
    def add_non_breaking_spaces(self, text: str) -> str:
        """
        Add non-breaking spaces to text for proper formatting.
        
        Legacy method - adds non-breaking spaces in specific contexts only.
        Preserved for backward compatibility with existing templates.
        
        For full GOST R 7.0.97-2016 compliance, use apply_full_gost_normalization().
        
        Note: OATI numbers (01-21-П-XXXXX/XX) should NOT have any spaces 
        between components as per user requirements. Removed space-adding
        regex to prevent unwanted spaces in document numbers.
        
        Args:
            text: Input text to process.
            
        Returns:
            Text with non-breaking spaces inserted where appropriate.
        """
        if not text:
            return text
        
        nbsp = '\u00A0'
        
        try:
            text = text.replace('№ ', f'№{nbsp}')
            text = text.replace('ФЗ «', f'ФЗ{nbsp}«')
            text = text.replace('№59-ФЗ', f'№{nbsp}59-ФЗ')
            text = text.replace('№ 59-ФЗ', f'№{nbsp}59-ФЗ')
            
            return text
            
        except Exception as e:
            logger.error(f"Error adding non-breaking spaces: {e}")
            return text
    
    def apply_full_gost_normalization(self, text: str) -> str:
        """
        Apply full GOST R 7.0.97-2016 text normalization.
        
        OPTIONAL method for comprehensive document text formatting according
        to Russian federal standards for organizational documentation.
        
        Applies:
        - Quotation marks normalization (« » instead of " ")
        - Dash types (hyphen, en dash, em dash) with proper spacing
        - Non-breaking spaces (after prepositions, abbreviations, №/ст./ч.)
        - Extra spaces cleanup
        
        WARNING: This is more aggressive than add_non_breaking_spaces().
        Only use when full GOST compliance is required and templates support it.
        
        Args:
            text: Input text to normalize
            
        Returns:
            Fully GOST-normalized text
            
        Examples:
            >>> gen = WordGenerator()
            >>> gen.apply_full_gost_normalization('В соответствии с "Законом" № 59-ФЗ')
            'В\xa0соответствии с\xa0«Законом» №\xa059-ФЗ'  # with NBSPs
        """
        if not text:
            return text
        
        try:
            return self.gost_normalizer.normalize_text(text)
        except Exception as e:
            logger.error(f"Error applying GOST normalization: {e}")
            return text  # Return original on error
    
    def _load_template_with_cache(self, template_path: str) -> Document:
        """
        Load template with caching for better batch processing performance.
        
        Caches template file content in memory and checks file modification time
        to detect changes. Creates a new Document from cached bytes on each call
        to ensure independent document instances.
        
        Args:
            template_path: Path to template file.
            
        Returns:
            New Document instance from template.
            
        Performance: ~10-50x faster for batch processing compared to loading from disk.
        """
        import io
        
        # Get file modification time
        file_mtime = os.path.getmtime(template_path)
        
        # Check if template is in cache and hasn't been modified
        if template_path in self._template_cache:
            cached_mtime, cached_bytes = self._template_cache[template_path]
            if cached_mtime == file_mtime:
                # Create new Document from cached bytes
                logger.debug(f"Loading template from cache: {template_path}")
                return Document(io.BytesIO(cached_bytes))
        
        # Load template from disk and cache it
        logger.debug(f"Loading template from disk: {template_path}")
        with open(template_path, 'rb') as f:
            template_bytes = f.read()
        
        # Store in cache
        self._template_cache[template_path] = (file_mtime, template_bytes)
        logger.debug(f"Template cached: {template_path} ({len(template_bytes)} bytes)")
        
        # Return new Document from bytes
        return Document(io.BytesIO(template_bytes))
    
    def replace_text_in_paragraph(
        self, 
        paragraph: Paragraph, 
        replacements: Dict[str, str]
    ) -> None:
        """
        Replace placeholder text in a document paragraph.
        
        Args:
            paragraph: Document paragraph object to process.
            replacements: Dictionary mapping placeholders to replacement values.
        """
        try:
            for key, value in replacements.items():
                if key in paragraph.text:
                    inline = paragraph.runs
                    for i in range(len(inline)):
                        if key in inline[i].text:
                            inline[i].text = inline[i].text.replace(key, value)
        except Exception as e:
            logger.error(f"Error replacing text in paragraph: {e}")
    
    def generate_document(
        self, 
        template_path: str, 
        data: Dict[str, str], 
        output_filename: str
    ) -> str:
        """
        Generate Word document from template with data replacement.
        
        Args:
            template_path: Path to template document file.
            data: Dictionary containing data for placeholder replacement.
            output_filename: Name for output document file.
            
        Returns:
            Full path to generated document.
            
        Raises:
            FileNotFoundError: If template file does not exist.
            Exception: If document generation fails.
        """
        logger.info(f"Generating document from template: {template_path}")
        
        if not os.path.exists(template_path):
            logger.error(f"Template file not found: {template_path}")
            raise FileNotFoundError(f"Template file not found: {template_path}")
        
        try:
            doc = self._load_template_with_cache(template_path)
            
            replacements = {
                '{{FULL_NAME_DATIVE}}': self.add_non_breaking_spaces(data.get('full_name_dative', '')),
                '{{SALUTATION}}': data.get('salutation', ''),
                '{{PORTAL_ID}}': data.get('portal_id', ''),
                '{{OATI_DATE}}': data.get('oati_date', ''),
                '{{OATI_NUMBER}}': self.add_non_breaking_spaces(data.get('oati_number', '')),
                '{{LAW_PART}}': data.get('law_part', ''),
                '{{DEPARTMENT}}': data.get('department', ''),
                '{{DEPARTMENTS}}': data.get('department', ''),
                '{{EMAIL}}': data.get('email', ''),
                '{{PORTAL_SOURCE}}': data.get('portal_source', ''),
                '{{UNION_SUFFIX}}': data.get('union_suffix', ''),
                '{{UNION_PARAGRAPH}}': data.get('union_paragraph', ''),
                '{{DEPARTMENTS_LIST}}': data.get('departments_list', ''),
            }
            
            for paragraph in doc.paragraphs:
                self.replace_text_in_paragraph(paragraph, replacements)
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self.replace_text_in_paragraph(paragraph, replacements)
            
            output_path = os.path.join(self.output_dir, output_filename)
            doc.save(output_path)
            
            # Verify file was actually created
            abs_output_path = os.path.abspath(output_path)
            if not os.path.exists(abs_output_path):
                raise IOError(f"Document was not saved successfully: {abs_output_path}")
            
            file_size = os.path.getsize(abs_output_path)
            if file_size == 0:
                raise IOError(f"Document was saved but is empty: {abs_output_path}")
            
            logger.info(f"Document generated successfully: {output_path} ({file_size} bytes)")
            return output_path
            
        except Exception as e:
            logger.error(f"Error generating document: {e}")
            raise
    
    def process_citizen_document(
        self, 
        citizen_data: Dict[str, Any], 
        declined_name: Dict[str, str], 
        salutation: str, 
        departments: List[str],
        portal_source: str = '', 
        union_suffix: str = '',
        union_paragraph: str = '', 
        departments_list: str = '',
        has_inspector: bool = False
    ) -> str:
        """
        Process citizen data and generate notification document.
        
        This is the main method for generating citizen notification documents.
        It orchestrates template selection, filename generation, and document creation.
        
        Args:
            citizen_data: Dictionary containing citizen information (name, OATI number, etc.).
            declined_name: Dictionary with declined name forms.
            salutation: Greeting text for the citizen.
            departments: List of involved department names.
            portal_source: Source portal identifier. Defaults to empty string.
            union_suffix: Suffix text for department unions. Defaults to empty string.
            union_paragraph: Additional paragraph for unions. Defaults to empty string.
            departments_list: Formatted list of departments. Defaults to empty string.
            has_inspector: Whether an inspector is involved (union_response flag). Defaults to False.
            
        Returns:
            Full path to generated document.
            
        Raises:
            KeyError: If required citizen data fields are missing.
            Exception: If document processing fails.
        """
        logger.info(f"Processing citizen document for: {citizen_data.get('last_name', 'Unknown')}")
        
        try:
            recipients = classify_recipients(departments)
            logger.info(f"Classified recipients: {recipients}")
            logger.debug(
                f"Recipient breakdown - Departments: {len(recipients.departments)}, "
                f"Prefectures: {len(recipients.prefectures)}, "
                f"OATI: {len(recipients.oati)}, "
                f"Combination key: '{recipients.get_combination_key()}'"
            )
            
            # CRITICAL: has_inspector flag indicates union_response (OATI handles case + forwards)
            # This determines template selection:
            # - has_inspector=True → template_ch4_oati_handles (with additional paragraph)
            # - has_inspector=False + multiple depts → template_ch4_multiple
            # - has_inspector=False + single dept → template_ch3
            law_part = citizen_data.get('law_part', '3')
            result = self.select_template(
                law_part, 
                departments, 
                has_inspector, 
                recipients_classification=recipients,
                union_response=has_inspector
            )
            
            # Handle both old (str) and new (tuple) return types for backward compatibility
            if isinstance(result, tuple):
                template_path, union_included, union_suffix_from_config = result
                logger.debug(
                    f"Received tuple from select_template: template_path='{template_path}', "
                    f"union_included={union_included}, union_suffix_from_config='{union_suffix_from_config}'"
                )
            else:
                # Fallback for backward compatibility if select_template returns just a string
                template_path = result
                union_included = False
                union_suffix_from_config = ""
                logger.warning("select_template returned string only (backward compatibility mode)")
            
            # Handle union_included flag: clear union data if not included
            if not union_included:
                logger.info("union_included=False: clearing union_suffix and union_paragraph")
                union_suffix = ""
                union_paragraph = ""
            elif not union_suffix:
                # Use from config if not explicitly provided
                union_suffix = union_suffix_from_config
                logger.debug(f"Using union_suffix from config: '{union_suffix}'")
            
            logger.info(f"Union configuration: union_included={union_included}, union_suffix='{union_suffix}'")
            
            first_initial = citizen_data['first_name'][0] if citizen_data.get('first_name') else ''
            middle_initial = citizen_data['middle_name'][0] if citizen_data.get('middle_name') else ''
            
            if 'last_name' not in citizen_data:
                raise KeyError("Missing required field: 'last_name' in citizen_data")
            
            output_filename = self.generate_filename(
                citizen_data['last_name'],
                first_initial,
                middle_initial,
                citizen_data.get('oati_number', '')
            )
            
            default_department = 'компетентный орган'
            
            document_data = {
                'full_name_dative': declined_name.get('full_name', ''),
                'salutation': salutation,
                'portal_id': citizen_data.get('portal_id', ''),
                'oati_date': citizen_data.get('oati_date', ''),
                'oati_number': citizen_data.get('oati_number', ''),
                'law_part': law_part,
                'department': departments_list if departments_list else default_department,
                'email': citizen_data.get('email', ''),
                'portal_source': portal_source,
                'union_suffix': union_suffix,
                'union_paragraph': union_paragraph,
                'departments_list': departments_list if departments_list else default_department,
            }
            
            output_path = self.generate_document(template_path, document_data, output_filename)
            
            # Detailed logging for debugging output issues
            abs_path = os.path.abspath(output_path)
            logger.info(f"Citizen document processed successfully: {output_path}")
            logger.info(f"Absolute path: {abs_path}")
            logger.info(f"File exists: {os.path.exists(abs_path)}")
            if os.path.exists(abs_path):
                file_size = os.path.getsize(abs_path)
                logger.info(f"File size: {file_size} bytes ({file_size/1024:.1f} KB)")
            
            return output_path
            
        except KeyError as e:
            logger.error(f"Missing required citizen data field: {e}")
            raise
        except Exception as e:
            logger.error(f"Error processing citizen document: {e}")
            raise


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    try:
        generator = WordGenerator()
        
        test_data = {
            'last_name': 'Иванов',
            'first_name': 'Иван',
            'middle_name': 'Петрович',
            'portal_id': '12345678',
            'oati_date': '01.01.2025',
            'oati_number': '01-21-П-1234/25',
            'law_part': '3',
            'email': 'test@example.com'
        }
        
        declined = {
            'full_name': 'Иванову Ивану Петровичу',
            'gender': 'male'
        }
        
        salutation = "Уважаемый Иван Петрович!"
        departments = ["Департамент тестирования"]
        
        result = generator.process_citizen_document(test_data, declined, salutation, departments)
        logger.info(f"Test document created: {result}")
        
    except Exception as e:
        logger.error(f"Test execution failed: {e}")
        raise
